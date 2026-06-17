"""桌面端主界面与业务流程编排模块。

本文件只负责把各个功能模块组织成可操作的 Tkinter 界面：
登录注册来自 auth.py，数据读写来自 db.py，检测推理来自 ai_core.py，
网络检测来自 network_utils.py，图表导出来自 visualization.py。

所有耗时操作都会通过 threading.Thread 放到后台执行，主线程只负责
界面绘制和控件刷新，因此摄像头检测、屏幕识别、批量处理和报告导出
不会把 GUI 卡死。
"""

import csv
import json
import shutil
import threading
import tkinter as tk
import tkinter.font as tkfont
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from queue import Empty, Full, Queue
from time import perf_counter
from tkinter import filedialog, messagebox, simpledialog, ttk

from ai_core import StudyBehaviorDetector
from auth import AuthError, AuthService
from db import Database
from network_utils import (
    DownloadError,
    NetworkError,
    check_url,
    check_urls_batch,
    download_file,
    download_image_from_url,
    test_connectivity,
)
from path_utils import resource_path, runtime_path
from visualization import (
    count_alert_labels,
    export_alert_chart,
    export_model_benchmark_chart,
    export_quantization_chart,
)


# 项目根路径统一从当前文件推导，打包后也能稳定定位模型和数据目录。
def write_runtime_error(context, exc_type, exc_value, exc_tb):
    """Write hidden GUI exceptions to a log file beside the exe or source tree."""
    log_dir = runtime_path("logs")
    log_dir.mkdir(parents=True, exist_ok=True)
    log_path = log_dir / "runtime_error.log"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    details = "".join(traceback.format_exception(exc_type, exc_value, exc_tb))
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"\n[{timestamp}] {context}\n")
        f.write(details)
        f.write("\n")
    return log_path


BASE_DIR = runtime_path(".")
MODELS_DIR = runtime_path("models")
MODEL_PT_PATH = resource_path("models/best.pt")
MODEL_FP16_PATH = MODELS_DIR / "best_fp16.pt"
MODEL_INT8_PATH = MODELS_DIR / "best_int8_quantized.pt"
MODEL_QUANT_REPORT_PATH = MODELS_DIR / "quantization_report.json"
MODEL_BENCHMARK_REPORT_PATH = MODELS_DIR / "benchmark_report.json"
RELABEL_DIR = BASE_DIR / "relabel"
RELABEL_IMAGES_DIR = RELABEL_DIR / "images"
RELABEL_LABELS_DIR = RELABEL_DIR / "labels"
SUSPECTED_DIR = RELABEL_DIR / "suspected_false_positive"
REPORTS_DIR = BASE_DIR / "reports"


# 统一维护界面色彩，方便把桌面程序做成接近网页应用的视觉风格。
COLORS = {
    "page": "#eef3f8",
    "card": "#ffffff",
    "sidebar": "#162033",
    "sidebar_hover": "#22304a",
    "accent": "#2563eb",
    "accent_hover": "#1d4ed8",
    "success": "#16a34a",
    "warning": "#f59e0b",
    "danger": "#dc2626",
    "text": "#111827",
    "muted": "#64748b",
    "border": "#d8e0ea",
    "soft": "#f8fafc",
}


# 常用字号集中定义，避免窗口不同区域字体大小不统一。
FONTS = {
    "hero": ("Segoe UI", 24, "bold"),
    "title": ("Segoe UI", 18, "bold"),
    "section": ("Segoe UI", 13, "bold"),
    "body": ("Segoe UI", 10),
    "small": ("Segoe UI", 9),
}


ACTION_TEXT = {
    "login": "用户登录",
    "register": "用户注册",
    "detect_image": "图片检测",
    "detect_batch": "批量检测",
    "start_camera": "启动摄像头检测",
    "stop_camera": "停止摄像头检测",
    "start_screen": "启动屏幕识别",
    "stop_screen": "停止屏幕识别",
    "check_url": "网络资源检测",
    "check_url_batch": "批量 URL 检测",
    "download_file": "下载文件测试",
    "download_image": "下载图片测试",
    "test_connectivity": "网络连通性测试",
    "crawl_images": "采集训练图片",
    "preprocess_dataset": "预处理图片集",
    "validate_dataset": "校验数据集",
    "generate_dataset_preview": "生成数据预览",
    "show_records": "查看检测记录",
    "show_logs": "查看操作日志",
    "export_chart": "导出统计图表",
    "export_logs": "导出日志 CSV",
    "generate_quantized_models": "生成量化模型",
    "export_quantization_chart": "导出量化对比图",
    "export_model_benchmark_chart": "导出模型性能图",
    "export_report": "导出分析报告",
    "save_suspected_false_positive": "保存疑似误报样本",
    "save_error_sample": "保存误识别样本",
    "record_behavior": "记录行为明细",
    "start_study_session": "开始学习会话",
    "end_study_session": "结束学习会话",
    "change_password": "修改密码",
    "reset_password": "重置用户密码",
    "change_role": "修改用户角色",
    "manage_users": "用户管理",
}


def configure_app_style(root):
    """配置 ttk 主题，让 tkinter 控件拥有更统一的现代视觉效果。"""
    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except tk.TclError:
        pass

    style.configure(".", font=FONTS["body"])
    style.configure("App.TFrame", background=COLORS["page"])
    style.configure("Card.TFrame", background=COLORS["card"])
    style.configure("Soft.TFrame", background=COLORS["soft"])
    style.configure("Muted.TLabel", background=COLORS["card"], foreground=COLORS["muted"])
    style.configure("Card.TLabel", background=COLORS["card"], foreground=COLORS["text"])
    style.configure("Page.TLabel", background=COLORS["page"], foreground=COLORS["text"])
    style.configure(
        "Hero.TLabel",
        background=COLORS["page"],
        foreground=COLORS["text"],
        font=FONTS["hero"],
    )
    style.configure(
        "Section.TLabel",
        background=COLORS["card"],
        foreground=COLORS["text"],
        font=FONTS["section"],
    )
    style.configure(
        "Accent.TButton",
        background=COLORS["accent"],
        foreground="#ffffff",
        borderwidth=0,
        focusthickness=0,
        padding=(16, 9),
    )
    style.map(
        "Accent.TButton",
        background=[("active", COLORS["accent_hover"]), ("disabled", "#aab7cc")],
        foreground=[("disabled", "#f8fafc")],
    )
    style.configure(
        "Secondary.TButton",
        background="#e8eef7",
        foreground=COLORS["text"],
        borderwidth=0,
        focusthickness=0,
        padding=(14, 8),
    )
    style.map("Secondary.TButton", background=[("active", "#dbe7f6")])
    style.configure(
        "Danger.TButton",
        background=COLORS["danger"],
        foreground="#ffffff",
        borderwidth=0,
        focusthickness=0,
        padding=(14, 8),
    )
    style.map("Danger.TButton", background=[("active", "#b91c1c")])
    style.configure(
        "Task.Horizontal.TProgressbar",
        background=COLORS["accent"],
        troughcolor="#dbe4ef",
        bordercolor=COLORS["border"],
        lightcolor=COLORS["accent"],
        darkcolor=COLORS["accent"],
    )
    style.configure(
        "Treeview",
        rowheight=30,
        borderwidth=0,
        background=COLORS["card"],
        fieldbackground=COLORS["card"],
        foreground=COLORS["text"],
    )
    style.configure(
        "Treeview.Heading",
        background=COLORS["soft"],
        foreground=COLORS["muted"],
        font=("Segoe UI", 9, "bold"),
        padding=(8, 6),
    )


def human_size(path):
    """把文件大小转换成 KB/MB 文本，便于模型对比窗口直接展示。"""
    path = Path(path)
    if not path.exists():
        return "未找到"
    size = path.stat().st_size
    units = ["B", "KB", "MB", "GB"]
    index = 0
    while size >= 1024 and index < len(units) - 1:
        size /= 1024
        index += 1
    return f"{size:.2f} {units[index]}"


def bytes_to_mb(size):
    """把字节数转换成 MB 浮点数，供量化对比图表使用。"""
    return round(size / (1024 * 1024), 4) if size else 0.0


def file_size_mb(path):
    """读取文件体积，文件不存在时返回 0，避免刷新表格时报错。"""
    path = Path(path)
    return bytes_to_mb(path.stat().st_size) if path.exists() else 0.0


def ratio_text(value):
    """把压缩比例转换成百分比文本，例如 0.5 显示为 50.0%。"""
    if value is None:
        return "-"
    return f"{value * 100:.1f}%"


def mean_value(values):
    """计算平均值，空列表返回 0，避免模型测试失败时除零。"""
    return sum(values) / len(values) if values else 0.0


def safe_alert_text(alerts_json):
    """把数据库中的告警 JSON 转成人能读懂的短文本。"""
    try:
        alerts = json.loads(alerts_json or "[]")
    except json.JSONDecodeError:
        return "解析失败"
    return ", ".join(alerts) if alerts else "无"


def short_path(value, max_len=56):
    """表格里展示路径时做中间省略，避免长路径把界面撑开。"""
    text = str(value or "-")
    if len(text) <= max_len:
        return text
    return f"{text[:24]}...{text[-24:]}"


def current_stamp():
    """生成适合文件名使用的时间戳。"""
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def is_image_file(path):
    """判断文件是否为常见图片格式，供批量检测任务分流。"""
    return Path(path).suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp"}


def is_video_file(path):
    """判断文件是否为常见视频格式，供批量检测任务分流。"""
    return Path(path).suffix.lower() in {".mp4", ".avi", ".mov", ".mkv", ".wmv"}


def labels_text_from_events(events):
    """把检测事件中的类别整理成展示文本。"""
    labels = sorted({event.get("label", "-") for event in events})
    return ", ".join(labels) if labels else "无"


class StatCard(tk.Frame):
    """仪表盘统计卡片，模拟网页中的信息卡样式。"""

    def __init__(self, master, title, value="0", note=""):
        super().__init__(
            master,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=18,
            pady=16,
        )
        self.value_var = tk.StringVar(value=value)
        self.note_var = tk.StringVar(value=note)

        tk.Label(
            self,
            text=title,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).pack(anchor="w")
        tk.Label(
            self,
            textvariable=self.value_var,
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Segoe UI", 21, "bold"),
        ).pack(anchor="w", pady=(8, 2))
        tk.Label(
            self,
            textvariable=self.note_var,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).pack(anchor="w")

    def update_value(self, value, note=None):
        """刷新卡片数字和说明文字。"""
        self.value_var.set(value)
        if note is not None:
            self.note_var.set(note)


class LoginFrame(ttk.Frame):
    """登录与注册界面，提供输入校验和友好的错误提示。"""

    def __init__(self, master, auth_service, on_login):
        super().__init__(master, style="App.TFrame")
        self.auth = auth_service
        self.on_login = on_login
        self.username_var = tk.StringVar()
        self.password_var = tk.StringVar()
        self.feedback_var = tk.StringVar(value="首次运行可使用 admin / admin123 登录。")
        self._build_layout()

    def _build_layout(self):
        """构建左右分栏登录页，让入口更接近现代网页应用。"""
        self.columnconfigure(0, weight=3)
        self.columnconfigure(1, weight=2)
        self.rowconfigure(0, weight=1)

        left = tk.Frame(self, bg=COLORS["sidebar"])
        left.grid(row=0, column=0, sticky="nsew")
        left.columnconfigure(0, weight=1)
        left.rowconfigure(0, weight=1)

        brand = tk.Frame(left, bg=COLORS["sidebar"], padx=64, pady=60)
        brand.grid(row=0, column=0, sticky="nsew")
        brand.rowconfigure(3, weight=1)
        tk.Label(
            brand,
            text="Study Behavior Monitor",
            bg=COLORS["sidebar"],
            fg="#ffffff",
            font=("Segoe UI", 28, "bold"),
            wraplength=420,
            justify="left",
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            brand,
            text="课堂学习状态智能识别系统",
            bg=COLORS["sidebar"],
            fg="#bfdbfe",
            font=("Microsoft YaHei UI", 16, "bold"),
        ).grid(row=1, column=0, sticky="w", pady=(16, 8))
        tk.Label(
            brand,
            text="支持图片检测、摄像头检测、网络资源检查、日志追踪和模型轻量化对比。",
            bg=COLORS["sidebar"],
            fg="#dbeafe",
            font=("Microsoft YaHei UI", 11),
            wraplength=460,
            justify="left",
        ).grid(row=2, column=0, sticky="w")

        right = tk.Frame(self, bg=COLORS["page"], padx=48, pady=48)
        right.grid(row=0, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(0, weight=1)

        card = tk.Frame(
            right,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=34,
            pady=32,
        )
        card.grid(row=0, column=0, sticky="nsew")
        card.columnconfigure(0, weight=1)

        tk.Label(
            card,
            text="欢迎登录",
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 22, "bold"),
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            card,
            text="请输入账号信息进入系统，注册普通用户后可直接登录。",
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["body"],
        ).grid(row=1, column=0, sticky="w", pady=(8, 24))

        self._build_form_field(card, "用户名", self.username_var, 2)
        self._build_form_field(card, "密码", self.password_var, 4, show="*")

        actions = tk.Frame(card, bg=COLORS["card"])
        actions.grid(row=6, column=0, sticky="ew", pady=(18, 0))
        actions.columnconfigure((0, 1), weight=1)
        ttk.Button(actions, text="登录", style="Accent.TButton", command=self.login).grid(
            row=0, column=0, sticky="ew", padx=(0, 8)
        )
        ttk.Button(
            actions,
            text="注册普通用户",
            style="Secondary.TButton",
            command=self.register,
        ).grid(row=0, column=1, sticky="ew", padx=(8, 0))

        tk.Label(
            card,
            textvariable=self.feedback_var,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
            wraplength=340,
            justify="left",
        ).grid(row=7, column=0, sticky="w", pady=(18, 0))

    def _build_form_field(self, parent, label, variable, row, show=None):
        """创建带标题的输入框，减少登录页重复布局代码。"""
        tk.Label(
            parent,
            text=label,
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 10, "bold"),
        ).grid(row=row, column=0, sticky="w", pady=(0, 6))
        entry = ttk.Entry(parent, textvariable=variable, show=show)
        entry.grid(row=row + 1, column=0, sticky="ew", ipady=6, pady=(0, 14))
        return entry

    def _read_credentials(self):
        """读取并校验登录表单，空输入会直接给出中文提示。"""
        username = self.username_var.get().strip()
        password = self.password_var.get()
        if not username or not password:
            self.feedback_var.set("请输入用户名和密码，两个字段都不能为空。")
            messagebox.showwarning("输入不完整", "请输入用户名和密码后再继续。")
            return None
        return username, password

    def login(self):
        """登录成功后进入主界面，失败时保留在当前页并提示原因。"""
        credentials = self._read_credentials()
        if credentials is None:
            return
        username, password = credentials
        try:
            user = self.auth.login(username, password)
        except AuthError as exc:
            self.feedback_var.set(str(exc))
            messagebox.showerror("登录失败", str(exc))
            return
        self.on_login(user)

    def register(self):
        """注册普通用户，账号创建成功后仍停留在登录页供用户确认登录。"""
        credentials = self._read_credentials()
        if credentials is None:
            return
        username, password = credentials
        try:
            user = self.auth.register(username, password)
        except AuthError as exc:
            self.feedback_var.set(str(exc))
            messagebox.showerror("注册失败", str(exc))
            return
        self.feedback_var.set(f"用户 {user['username']} 已创建，可以直接登录。")
        messagebox.showinfo("注册成功", f"用户 {user['username']} 已创建。")


class MainFrame(ttk.Frame):
    """主功能界面，集中管理按钮入口、后台任务、状态栏和数据展示。"""

    def __init__(self, master, database, auth_service, user):
        super().__init__(master, style="App.TFrame")
        self.db = database
        self.auth = auth_service
        self.user = user
        self.status_var = tk.StringVar(value="Ready")
        self.task_count_var = tk.StringVar(value="0")
        self.camera_state_var = tk.StringVar(value="摄像头：未运行")
        self.running = False
        self.active_tasks = 0
        self.camera_stop_event = None
        self.last_detection = None
        self.camera_queue = Queue(maxsize=2)
        self.camera_preview_label = None
        self.camera_stats_var = tk.StringVar(value="摄像头未启动。")
        self.camera_image_ref = None
        self.camera_latest_frame = None
        self.camera_latest_events = []
        self.camera_latest_summary = {}
        self.camera_session = None
        self.screen_running = False
        self.screen_stop_event = None
        self.screen_control_window = None
        self.screen_control_status_var = tk.StringVar(value="等待屏幕识别启动。")
        self.screen_session = None
        self.model_benchmark_metrics = []
        self.batch_status_var = tk.StringVar(value="暂无批量任务。")
        self.sidebar_visible = True
        self.workspace_zoom = 1.0
        self.sidebar = None
        self.sidebar_rail = None
        self.main_canvas = None
        self.content_frame = None
        self.active_page = "home"
        self.camera_start_button = None
        self.camera_stop_button = None
        self.screen_start_button = None
        self.screen_stop_button = None
        self.live_sample_button = None
        self.nav_buttons = {}
        self.stat_cards = {}
        self.records_tree = None
        self.logs_tree = None
        self.save_sample_button = None
        self._build_layout()
        self.refresh_dashboard()

    def _db_supports(self, method_name):
        """检查数据库模块是否提供增强接口，兼容旧数据库文件和新 PR 版本。"""
        return callable(getattr(self.db, method_name, None))

    def _behavior_type_from_label(self, label):
        """把模型类别映射到数据库约束允许的行为类型。"""
        normalized = str(label or "normal").lower()
        return normalized if normalized in {"normal", "phone", "sleep", "eat"} else "other"

    def _record_behavior_events(self, events, source_type, source_path, output_image_path=""):
        """把一次检测的事件写入行为明细表，供后续统计图和报告复用。"""
        if not self._db_supports("add_behavior_record"):
            return
        records = events or [
            {"label": "normal", "confidence": 1.0, "alert": False, "reason": "未检测到异常行为。"}
        ]
        db_source_type = "video" if source_type == "screen" else source_type
        for event in records:
            extra_info = dict(event)
            extra_info["original_source_type"] = source_type
            try:
                self.db.add_behavior_record(
                    self.user["id"],
                    self._behavior_type_from_label(event.get("label")),
                    float(event.get("confidence", 1.0)),
                    session_id=event.get("session_id"),
                    is_alert=bool(event.get("alert", False)),
                    alert_reason=event.get("reason", ""),
                    source_type=db_source_type,
                    source_path=str(source_path or ""),
                    image_path=str(source_path or ""),
                    output_image_path=str(output_image_path or ""),
                    extra_info=extra_info,
                )
            except Exception as exc:
                self.db.log_operation(self.user["id"], "record_behavior_failed", str(exc))
                break

    def _attach_study_session(self, session, session_name):
        """为摄像头或屏幕识别建立学习会话，后续结束时可统计时长和告警。"""
        if not session or not self._db_supports("start_study_session"):
            return
        try:
            session["study_session_id"] = self.db.start_study_session(self.user["id"], session_name)
            self.db.log_operation(self.user["id"], "start_study_session", session_name)
        except Exception as exc:
            session["study_session_id"] = None
            self.db.log_operation(self.user["id"], "start_study_session_failed", str(exc))

    def _finish_study_session(self, session, notes=""):
        """结束学习会话，并把聚合后的告警统计写入数据库。"""
        if not session or not self._db_supports("end_study_session") or not session.get("study_session_id"):
            return
        label_counts = session.get("label_counts") or {}
        alerts = session.get("alerts") or set()
        alert_total = sum(label_counts.get(label, 0) for label in alerts)
        alert_stats = {
            "total": alert_total,
            "phone": label_counts.get("phone", 0),
            "sleep": label_counts.get("sleep", 0),
            "eat": label_counts.get("eat", 0),
        }
        try:
            self.db.end_study_session(
                session["study_session_id"],
                focus_score=self._estimate_focus_score(session),
                effective_seconds=self._estimate_effective_seconds(session),
                notes=notes,
                alert_stats=alert_stats,
            )
            self.db.log_operation(self.user["id"], "end_study_session", notes)
        except Exception as exc:
            self.db.log_operation(self.user["id"], "end_study_session_failed", str(exc))

    def _estimate_focus_score(self, session):
        """根据告警密度估算一个 0-100 的专注分数，供会话报表初步展示。"""
        frames = max(1, int(session.get("frame_count", 0)))
        alerts = len(session.get("alerts") or [])
        event_density = min(1.0, float(session.get("event_count", 0)) / frames)
        return max(0, min(100, int(100 - alerts * 12 - event_density * 18)))

    def _estimate_effective_seconds(self, session):
        """用会话总时长减去告警占比粗略估算有效学习时长。"""
        started_at = session.get("started_at") or datetime.now()
        total_seconds = max(0, int((datetime.now() - started_at).total_seconds()))
        alerts = len(session.get("alerts") or [])
        penalty = min(total_seconds, alerts * 30)
        return max(0, total_seconds - penalty)

    def _count_db_table(self, table_name):
        """统计指定表的行数，表不存在时返回 0，适配渐进式数据库升级。"""
        if not table_name.replace("_", "").isalnum():
            return 0
        try:
            with self.db.connect() as conn:
                row = conn.execute(f"SELECT COUNT(*) AS total FROM {table_name}").fetchone()
                return int(row["total"]) if row else 0
        except Exception:
            return 0

    def _alert_stat_records(self, limit=1000):
        """优先把行为明细转换成图表统计输入，没有明细时退回检测记录。"""
        if self._db_supports("get_behavior_records"):
            try:
                behavior_rows = self.db.get_behavior_records(limit=limit)
            except Exception:
                behavior_rows = []
            if behavior_rows:
                return [
                    {"alerts_json": json.dumps([row["behavior_type"]], ensure_ascii=False)}
                    for row in behavior_rows
                    if int(row.get("is_alert") or 0)
                ]
        return self.db.list_detection_records(limit=limit)

    def _build_layout(self):
        """搭建侧边栏和仪表盘，使桌面程序呈现网页式信息架构。"""
        self.columnconfigure(1, weight=1)
        self.rowconfigure(0, weight=1)

        self.sidebar_rail = tk.Frame(self, bg=COLORS["sidebar"], width=54, padx=8, pady=16)
        self.sidebar_rail.grid(row=0, column=0, sticky="ns")
        self.sidebar_rail.grid_propagate(False)
        ttk.Button(
            self.sidebar_rail,
            text="☰",
            style="Secondary.TButton",
            command=self.toggle_sidebar,
        ).pack(fill="x")
        self.sidebar_rail.grid_remove()

        self.sidebar = tk.Frame(self, bg=COLORS["sidebar"], width=250, padx=18, pady=20)
        self.sidebar.grid(row=0, column=0, sticky="ns")
        self.sidebar.grid_propagate(False)
        self.sidebar.columnconfigure(0, weight=1)

        ttk.Button(
            self.sidebar,
            text="隐藏工具栏",
            style="Secondary.TButton",
            command=self.toggle_sidebar,
        ).grid(row=0, column=0, sticky="ew", pady=(0, 12))
        tk.Label(
            self.sidebar,
            text="SBM",
            bg=COLORS["sidebar"],
            fg="#ffffff",
            font=("Segoe UI", 24, "bold"),
        ).grid(row=1, column=0, sticky="w")
        tk.Label(
            self.sidebar,
            text=f"{self.user['username']} · {self.user['role']}",
            bg=COLORS["sidebar"],
            fg="#bfdbfe",
            font=FONTS["small"],
        ).grid(row=2, column=0, sticky="w", pady=(2, 18))

        nav_groups = [
            (
                "检测",
                [
                    ("图片检测", self.detect_image),
                    ("批量图片/视频", self.detect_batch_files),
                    ("误报规则", self.show_false_positive_rules),
                ],
            ),
            (
                "记录",
                [
                    ("检测记录", self.show_records),
                    ("操作日志", self.show_logs),
                ],
            ),
            (
                "网络监测",
                [
                    ("网络连通性", self.check_network_connectivity),
                    ("URL 可达性检测", self.check_model_url),
                    ("批量 URL 检测", self.check_url_batch),
                    ("下载文件测试", self.download_remote_file),
                    ("下载图片测试", self.download_remote_image),
                ],
            ),
            (
                "数据库查询",
                [
                    ("数据库概览", self.show_database_summary),
                    ("误报样本", self.show_misclassified_samples),
                    ("学习会话", self.show_study_sessions),
                    ("行为明细", self.show_behavior_records),
                    ("训练周期", self.show_training_cycles),
                ],
            ),
            (
                "数据工具",
                [
                    ("采集训练图片", self.crawl_training_images),
                    ("预处理图片集", self.run_preprocess_pipeline),
                    ("校验数据集", self.validate_training_dataset),
                    ("生成数据预览", self.generate_dataset_preview),
                ],
            ),
            (
                "结果导出",
                [
                    ("导出图表", self.export_chart),
                    ("导出报告", self.export_analysis_report),
                    ("模型对比", self.show_model_compare),
                ],
            ),
            (
                "账号管理",
                [
                    ("修改密码", self.change_password),
                    ("用户管理", self.manage_users if self.auth.is_admin(self.user) else self._admin_only_notice),
                ],
            ),
        ]
        for row, (label, items) in enumerate(nav_groups, start=3):
            self._build_nav_menu(self.sidebar, row, label, items)

        self.save_sample_button = self.nav_buttons.get("误识别样本")
        self._set_camera_controls(running=False)
        self._set_screen_controls(running=False)

        main_area = tk.Frame(self, bg=COLORS["page"])
        main_area.grid(row=0, column=1, sticky="nsew")
        main_area.columnconfigure(0, weight=1)
        main_area.rowconfigure(1, weight=1)

        top_tools = tk.Frame(main_area, bg=COLORS["page"], padx=14, pady=10)
        top_tools.grid(row=0, column=0, sticky="ew", pady=(0, 0))
        top_tools.columnconfigure(1, weight=1)
        ttk.Button(
            top_tools,
            text="首页",
            style="Secondary.TButton",
            command=self.show_home,
        ).grid(row=0, column=0, sticky="w")
        ttk.Button(
            top_tools,
            text="缩小",
            style="Secondary.TButton",
            command=lambda: self.adjust_workspace_zoom(-0.1),
        ).grid(row=0, column=2, padx=(8, 0))
        ttk.Button(
            top_tools,
            text="放大",
            style="Secondary.TButton",
            command=lambda: self.adjust_workspace_zoom(0.1),
        ).grid(row=0, column=3, padx=(8, 0))

        self.main_canvas = tk.Canvas(main_area, bg=COLORS["page"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_area, orient="vertical", command=self.main_canvas.yview)
        self.main_canvas.configure(yscrollcommand=scrollbar.set)
        self.main_canvas.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")

        content = tk.Frame(self.main_canvas, bg=COLORS["page"], padx=26, pady=22)
        self.content_frame = content
        canvas_window = self.main_canvas.create_window((0, 0), window=content, anchor="nw")
        content.bind(
            "<Configure>",
            lambda event: self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all")),
        )
        self.main_canvas.bind(
            "<Configure>",
            lambda event: self.main_canvas.itemconfigure(canvas_window, width=event.width),
        )
        self.main_canvas.bind_all("<MouseWheel>", self._on_workspace_mousewheel)
        self.main_canvas.bind_all("<Control-MouseWheel>", self._on_zoom_mousewheel)
        content.columnconfigure(0, weight=1)
        content.rowconfigure(4, weight=1)

        self.show_home()
        self.after(120, self._poll_camera_frames)

    def show_home(self):
        """回到首页仪表盘，把右侧主台恢复为工作台首页。"""
        self.active_page = "home"
        self._clear_main_content()
        content = self.content_frame
        content.columnconfigure(0, weight=1)
        content.rowconfigure(4, weight=1)
        self._build_header(content)
        self._build_stats(content)
        self._build_action_cards(content)
        self._build_camera_panel(content)
        self._build_recent_tables(content)
        self._build_status_bar(content)
        self.apply_visual_zoom()
        self.after(0, self._sync_workspace_layout)
        self.refresh_dashboard()

    def _clear_main_content(self):
        """清空右侧主台内容，用于在首页和工具页面之间切换。"""
        if self.content_frame is None:
            return
        for child in self.content_frame.winfo_children():
            child.destroy()
        for index in range(8):
            self.content_frame.rowconfigure(index, weight=0)
            self.content_frame.columnconfigure(index, weight=0)
        self.records_tree = None
        self.logs_tree = None
        self.stat_cards = {}
        self.tool_progress = {}
        self._reset_workspace_scroll()

    def _reset_workspace_scroll(self):
        """Reset the right workspace canvas after page switches."""
        if self.main_canvas is None:
            return
        try:
            self.main_canvas.update_idletasks()
            self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))
            self.main_canvas.yview_moveto(0)
        except tk.TclError:
            pass

    def _sync_workspace_layout(self):
        """Synchronize canvas scrollregion after Tk finishes creating widgets."""
        self._reset_workspace_scroll()

    def _open_main_page(self, title, show_title=True):
        """在右侧主台打开完整页面，代替小弹窗。"""
        self.active_page = title
        self._clear_main_content()
        page = tk.Frame(self.content_frame, bg=COLORS["page"])
        page.grid(row=0, column=0, sticky="nsew")
        page.columnconfigure(0, weight=1)
        page.rowconfigure(1, weight=1)
        self.content_frame.columnconfigure(0, weight=1)
        self.content_frame.rowconfigure(0, weight=1)
        if show_title:
            tk.Label(
                page,
                text=title,
                bg=COLORS["page"],
                fg=COLORS["text"],
                font=("Microsoft YaHei UI", 22, "bold"),
            ).grid(row=0, column=0, sticky="w", pady=(0, 14))
        self.apply_visual_zoom()
        self.after(0, self._sync_workspace_layout)
        return page

    def _build_nav_menu(self, parent, row, label, items):
        """创建左侧分组入口，点击后在右侧主台打开完整工具页。"""
        nav_button = ttk.Button(
            parent,
            text=label,
            style="Secondary.TButton",
            command=lambda title=label, actions=items: self._show_tool_group_page(title, actions),
        )
        nav_button.grid(row=row, column=0, sticky="ew", pady=4)
        for item_label, _ in items:
            self.nav_buttons[item_label] = nav_button

    def _show_tool_group_page(self, title, items):
        """把左侧工具分组展示成完整页面，避免用户在小下拉菜单里寻找功能。"""
        descriptions = {
            "检测": "集中运行图片检测、批量处理，并查看误报样本的保存规则。",
            "记录": "查看系统产生的检测记录和操作日志，便于复盘和验收。",
            "网络监测": "检查远程资源地址是否可达，网络数据库功能可在后续继续接入。",
            "数据库查询": "查看本地 SQLite 数据库存储概况，确认记录是否写入。",
            "结果导出": "导出统计图、分析报告，并进入模型量化性能对比页面。",
            "账号管理": "修改当前账号密码，管理员可查看和维护用户列表。",
        }
        page = self._open_main_page(title)
        page.rowconfigure(1, weight=1)
        tk.Label(
            page,
            text=descriptions.get(title, "选择下方功能继续操作。"),
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=FONTS["body"],
            wraplength=980,
            justify="left",
        ).grid(row=1, column=0, sticky="ew", pady=(0, 18))

        action_grid = tk.Frame(page, bg=COLORS["page"])
        action_grid.grid(row=2, column=0, sticky="nsew")
        for column in range(3):
            action_grid.columnconfigure(column, weight=1)
        for index, (item_label, command) in enumerate(items):
            row, column = divmod(index, 3)
            card = tk.Frame(
                action_grid,
                bg=COLORS["card"],
                highlightbackground=COLORS["border"],
                highlightthickness=1,
                padx=18,
                pady=16,
            )
            card.grid(row=row, column=column, sticky="nsew", padx=(0 if column == 0 else 12, 0), pady=(0, 12))
            card.columnconfigure(0, weight=1)
            tk.Label(
                card,
                text=item_label,
                bg=COLORS["card"],
                fg=COLORS["text"],
                font=("Microsoft YaHei UI", 13, "bold"),
            ).grid(row=0, column=0, sticky="w")
            button = ttk.Button(
                card,
                text="打开",
                style="Accent.TButton",
                command=lambda action=command: action(),
            )
            button.grid(row=1, column=0, sticky="ew", pady=(14, 0))
            if title == "数据工具":
                self._attach_tool_progress(card, item_label, button)
        self.apply_visual_zoom()
        self.after(0, self._sync_workspace_layout)

    def _attach_tool_progress(self, card, item_label, button):
        """Show the current data-tool progress inside its card."""
        progress_var = tk.DoubleVar(value=0)
        percent_var = tk.StringVar(value="0%")
        status_var = tk.StringVar(value="待开始")
        row = tk.Frame(card, bg=COLORS["card"])
        row.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        row.columnconfigure(0, weight=1)
        bar = ttk.Progressbar(
            row,
            variable=progress_var,
            maximum=100,
            mode="determinate",
            style="Task.Horizontal.TProgressbar",
        )
        bar.grid(row=0, column=0, sticky="ew")
        tk.Label(
            row,
            textvariable=percent_var,
            bg=COLORS["card"],
            fg=COLORS["accent"],
            font=("Microsoft YaHei UI", 10, "bold"),
            width=5,
            anchor="e",
        ).grid(row=0, column=1, sticky="e", padx=(10, 0))
        tk.Label(
            card,
            textvariable=status_var,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).grid(row=3, column=0, sticky="w", pady=(6, 0))
        self.tool_progress[item_label] = {
            "progress": progress_var,
            "percent": percent_var,
            "status": status_var,
            "button": button,
        }

    def _reset_tool_progress(self, item_label, text="准备开始..."):
        state = self.tool_progress.get(item_label)
        if not state:
            return
        state["progress"].set(0)
        state["percent"].set("0%")
        state["status"].set(text)
        state["button"].configure(state="disabled")

    def _post_tool_progress(self, item_label, done, total, text):
        percent = 100 if total <= 0 else min(100, max(0, done / total * 100))

        def update():
            state = self.tool_progress.get(item_label)
            if not state:
                return
            state["progress"].set(percent)
            state["percent"].set(f"{percent:.0f}%")
            state["status"].set(f"{text}（{percent:.0f}%）")

        self.after(0, update)

    def _finish_tool_progress(self, item_label, text="已完成"):
        def update():
            state = self.tool_progress.get(item_label)
            if not state:
                return
            state["progress"].set(100)
            state["percent"].set("100%")
            state["status"].set(text)
            state["button"].configure(state="normal")

        self.after(0, update)

    def _fail_tool_progress(self, item_label, text="执行失败"):
        def update():
            state = self.tool_progress.get(item_label)
            if not state:
                return
            state["status"].set(text)
            state["button"].configure(state="normal")

        self.after(0, update)

    def toggle_sidebar(self):
        """折叠或展开左侧工具栏，为右侧主台留出更多空间。"""
        if self.sidebar_visible:
            self.sidebar.grid_remove()
            self.sidebar_rail.grid()
        else:
            self.sidebar_rail.grid_remove()
            self.sidebar.grid()
        self.sidebar_visible = not self.sidebar_visible

    def adjust_workspace_zoom(self, delta):
        """调整主台视觉字体大小，不改变数据和检测结果本身。"""
        self.workspace_zoom = min(1.35, max(0.8, self.workspace_zoom + delta))
        self.apply_visual_zoom()

    def _on_workspace_mousewheel(self, event):
        """让右侧主台支持鼠标滚轮上下滑动。"""
        if event.state & 0x0004:
            self._on_zoom_mousewheel(event)
            return "break"
        if self.main_canvas is not None:
            self.main_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_zoom_mousewheel(self, event):
        """支持 Ctrl+滚轮或触控板捏合映射事件进行视觉字体缩放。"""
        self.adjust_workspace_zoom(0.05 if event.delta > 0 else -0.05)
        return "break"

    def apply_visual_zoom(self):
        """按当前缩放比例刷新 ttk 样式和已创建控件字体。"""
        size = max(8, int(10 * self.workspace_zoom))
        heading_size = max(10, int(13 * self.workspace_zoom))
        title_size = max(14, int(18 * self.workspace_zoom))
        style = ttk.Style(self)
        style.configure(".", font=("Segoe UI", size))
        style.configure("Treeview", font=("Segoe UI", size), rowheight=max(26, int(30 * self.workspace_zoom)))
        style.configure("Treeview.Heading", font=("Segoe UI", max(8, int(9 * self.workspace_zoom)), "bold"))
        style.configure("Section.TLabel", font=("Segoe UI", heading_size, "bold"))
        style.configure("Hero.TLabel", font=("Segoe UI", title_size, "bold"))
        for root in [self.sidebar, self.sidebar_rail, self.content_frame]:
            if root is not None and root.winfo_exists():
                self._scale_widget_tree(root)

    def _scale_widget_tree(self, widget):
        """递归缩放 tk 控件字体；ttk 控件主要由 style 控制。"""
        try:
            font_value = widget.cget("font")
        except tk.TclError:
            font_value = None
        if font_value:
            if not hasattr(widget, "_base_font_actual"):
                widget._base_font_actual = tkfont.Font(font=font_value).actual()
            actual = widget._base_font_actual
            base_size = abs(actual.get("size", 10)) or 10
            scaled_size = max(8, int(base_size * self.workspace_zoom))
            weight = actual.get("weight", "normal")
            slant = actual.get("slant", "roman")
            widget.configure(font=(actual.get("family", "Segoe UI"), scaled_size, weight, slant))
        for child in widget.winfo_children():
            self._scale_widget_tree(child)

    def _admin_only_notice(self):
        """普通用户点击管理员功能时给出权限提示。"""
        messagebox.showinfo("权限不足", "该功能需要管理员账号。")

    def _build_header(self, parent):
        """构建顶部欢迎区，提供当前身份和核心状态。"""
        header = tk.Frame(parent, bg=COLORS["page"])
        header.grid(row=0, column=0, sticky="ew")
        header.columnconfigure(0, weight=1)

        tk.Label(
            header,
            text="学习行为监测工作台",
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 24, "bold"),
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            header,
            text="面向课堂实验的图片识别、摄像头检测、日志追踪和模型部署辅助工具。",
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=("Microsoft YaHei UI", 10),
        ).grid(row=1, column=0, sticky="w", pady=(6, 0))

        status_pill = tk.Frame(
            header,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=16,
            pady=10,
        )
        status_pill.grid(row=0, column=1, rowspan=2, sticky="e")
        tk.Label(
            status_pill,
            text="当前状态",
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).pack(anchor="e")
        tk.Label(
            status_pill,
            textvariable=self.status_var,
            bg=COLORS["card"],
            fg=COLORS["accent"],
            font=("Segoe UI", 11, "bold"),
        ).pack(anchor="e", pady=(2, 0))

    def _build_stats(self, parent):
        """构建四个统计卡片，用于快速了解系统运行情况。"""
        stats = tk.Frame(parent, bg=COLORS["page"])
        stats.grid(row=1, column=0, sticky="ew", pady=(22, 16))
        for column in range(4):
            stats.columnconfigure(column, weight=1)

        cards = [
            ("records", "检测记录", "0", "最近检测结果"),
            ("logs", "操作日志", "0", "最近操作行为"),
            ("model", "模型文件", human_size(MODEL_PT_PATH), "PyTorch 权重"),
            ("tasks", "后台任务", "0", "线程并发执行"),
        ]
        for column, (key, title, value, note) in enumerate(cards):
            card = StatCard(stats, title, value, note)
            card.grid(row=0, column=column, sticky="ew", padx=(0 if column == 0 else 10, 0))
            self.stat_cards[key] = card

    def _build_action_cards(self, parent):
        """构建常用操作卡片，减少用户在侧边栏和弹窗之间来回寻找。"""
        actions = tk.Frame(parent, bg=COLORS["page"])
        actions.grid(row=2, column=0, sticky="ew", pady=(0, 18))
        for column in range(4):
            actions.columnconfigure(column, weight=1)

        items = [
            ("图片检测", "选择本地图片并输出标注图。", "开始检测", self.detect_image),
            ("屏幕识别", "检测线上课堂或会议窗口画面。", "开始识别", self.start_screen_detection),
            ("批量检测", "多图片并发、视频队列化处理。", "选择文件", self.detect_batch_files),
            ("分析报告", "导出 Word 和 PDF 课堂分析报告。", "导出报告", self.export_analysis_report),
        ]
        for column, item in enumerate(items):
            self._build_action_card(actions, column, *item)

    def _build_action_card(self, parent, column, title, desc, button_text, command):
        """创建单个操作卡片，保持按钮宽度和卡片间距统一。"""
        card = tk.Frame(
            parent,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=16,
            pady=14,
        )
        card.grid(row=0, column=column, sticky="nsew", padx=(0 if column == 0 else 10, 0))
        card.columnconfigure(0, weight=1)
        tk.Label(
            card,
            text=title,
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 12, "bold"),
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            card,
            text=desc,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
            wraplength=180,
            justify="left",
        ).grid(row=1, column=0, sticky="w", pady=(6, 12))
        ttk.Button(
            card,
            text=button_text,
            style="Secondary.TButton",
            command=command,
        ).grid(row=2, column=0, sticky="ew")

    def _build_camera_panel(self, parent):
        """在主界面内嵌摄像头画面，替代 OpenCV 弹出窗口。"""
        panel = tk.Frame(
            parent,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=16,
            pady=14,
        )
        panel.grid(row=3, column=0, sticky="ew", pady=(0, 18))
        panel.columnconfigure(0, weight=3)
        panel.columnconfigure(1, weight=1)

        preview = tk.Frame(panel, bg="#0f172a", width=720, height=360)
        preview.grid(row=0, column=0, sticky="nsew", padx=(0, 16))
        preview.grid_propagate(False)
        preview.columnconfigure(0, weight=1)
        preview.rowconfigure(0, weight=1)

        self.camera_preview_label = tk.Label(
            preview,
            text="摄像头画面将在这里显示",
            bg="#0f172a",
            fg="#cbd5e1",
            font=("Microsoft YaHei UI", 12, "bold"),
        )
        self.camera_preview_label.grid(row=0, column=0, sticky="nsew")

        info = tk.Frame(panel, bg=COLORS["card"])
        info.grid(row=0, column=1, sticky="nsew")
        info.columnconfigure(0, weight=1)
        tk.Label(
            info,
            text="实时摄像头检测",
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 13, "bold"),
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            info,
            textvariable=self.camera_stats_var,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["body"],
            wraplength=300,
            justify="left",
        ).grid(row=1, column=0, sticky="nw", pady=(10, 16))

        buttons = tk.Frame(info, bg=COLORS["card"])
        buttons.grid(row=2, column=0, sticky="ew")
        buttons.columnconfigure((0, 1), weight=1)
        self.camera_start_button = ttk.Button(
            buttons,
            text="启动",
            style="Accent.TButton",
            command=self.start_camera,
        )
        self.camera_start_button.grid(row=0, column=0, sticky="ew", padx=(0, 8))
        self.camera_stop_button = ttk.Button(
            buttons,
            text="停止",
            style="Danger.TButton",
            command=self.stop_camera,
        )
        self.camera_stop_button.grid(row=0, column=1, sticky="ew")
        self.screen_start_button = ttk.Button(
            buttons,
            text="屏幕识别",
            style="Secondary.TButton",
            command=self.start_screen_detection,
        )
        self.screen_start_button.grid(row=1, column=0, sticky="ew", padx=(0, 8), pady=(8, 0))
        self.screen_stop_button = ttk.Button(
            buttons,
            text="停止屏幕",
            style="Danger.TButton",
            command=self.stop_screen_detection,
        )
        self.screen_stop_button.grid(row=1, column=1, sticky="ew", pady=(8, 0))
        self.live_sample_button = ttk.Button(
            buttons,
            text="保存误报帧",
            style="Secondary.TButton",
            command=lambda: self._save_current_live_frame("live_manual"),
        )
        self.live_sample_button.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        self._set_live_sample_button_enabled(self.camera_latest_frame is not None)
        tk.Label(
            info,
            textvariable=self.batch_status_var,
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
            wraplength=300,
            justify="left",
        ).grid(row=3, column=0, sticky="nw", pady=(14, 0))

    def _build_recent_tables(self, parent):
        """首页展示最近记录和日志，减少频繁打开弹窗的成本。"""
        panel = tk.Frame(
            parent,
            bg=COLORS["card"],
            highlightbackground=COLORS["border"],
            highlightthickness=1,
            padx=16,
            pady=14,
        )
        panel.grid(row=4, column=0, sticky="nsew")
        panel.columnconfigure(0, weight=1)
        panel.rowconfigure(1, weight=1)

        title_bar = tk.Frame(panel, bg=COLORS["card"])
        title_bar.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        title_bar.columnconfigure(0, weight=1)
        tk.Label(
            title_bar,
            text="最近活动",
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 13, "bold"),
        ).grid(row=0, column=0, sticky="w")
        ttk.Button(
            title_bar,
            text="刷新",
            style="Secondary.TButton",
            command=self.refresh_dashboard,
        ).grid(row=0, column=1, sticky="e")

        notebook = ttk.Notebook(panel)
        notebook.grid(row=1, column=0, sticky="nsew")

        records_frame = ttk.Frame(notebook, style="Card.TFrame")
        logs_frame = ttk.Frame(notebook, style="Card.TFrame")
        records_frame.columnconfigure(0, weight=1)
        records_frame.rowconfigure(0, weight=1)
        logs_frame.columnconfigure(0, weight=1)
        logs_frame.rowconfigure(0, weight=1)
        notebook.add(records_frame, text="检测记录")
        notebook.add(logs_frame, text="操作日志")

        self.records_tree = self._create_tree(
            records_frame,
            [
                ("time", "时间", 150),
                ("user", "用户", 90),
                ("alerts", "告警", 120),
                ("source", "来源", 280),
            ],
        )
        self.logs_tree = self._create_tree(
            logs_frame,
            [
                ("time", "时间", 150),
                ("user", "用户", 90),
                ("action", "操作", 120),
                ("detail", "详情", 360),
            ],
        )

    def _build_status_bar(self, parent):
        """底部状态栏展示线程数量和摄像头运行状态。"""
        bar = tk.Frame(parent, bg=COLORS["page"])
        bar.grid(row=5, column=0, sticky="ew", pady=(12, 0))
        bar.columnconfigure(0, weight=1)
        tk.Label(
            bar,
            textvariable=self.camera_state_var,
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            bar,
            textvariable=self.task_count_var,
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=FONTS["small"],
        ).grid(row=0, column=1, sticky="e")

    def _create_tree(self, parent, columns):
        """创建带滚动条的表格控件，统一首页和弹窗的表格风格。"""
        tree = ttk.Treeview(parent, columns=[item[0] for item in columns], show="headings")
        for key, title, width in columns:
            tree.heading(key, text=title)
            tree.column(key, width=width, anchor="w", stretch=True)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        tree.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        return tree

    def refresh_dashboard(self):
        """刷新首页统计卡片、近期检测记录和近期日志。"""
        if self.active_page != "home" or not self.stat_cards or self.records_tree is None:
            return
        records = self.db.list_detection_records(limit=20)
        logs = self.db.list_operation_logs(limit=20)
        quant_note = f"FP16: {human_size(MODEL_FP16_PATH)} / INT8: {human_size(MODEL_INT8_PATH)}"

        self.stat_cards["records"].update_value(str(len(records)), "最近 20 条以内")
        self.stat_cards["logs"].update_value(str(len(logs)), "最近 20 条以内")
        self.stat_cards["model"].update_value(human_size(MODEL_PT_PATH), quant_note)
        self.stat_cards["tasks"].update_value(str(self.active_tasks), "线程并发执行")
        self.task_count_var.set(f"后台任务：{self.active_tasks}")
        self._fill_records_tree(records)
        self._fill_logs_tree(logs)

    def _fill_records_tree(self, records):
        """把检测记录写入首页表格。"""
        self.records_tree.delete(*self.records_tree.get_children())
        for record in records:
            self.records_tree.insert(
                "",
                "end",
                values=(
                    record["created_at"],
                    record["username"] or "-",
                    safe_alert_text(record["alerts_json"]),
                    short_path(record["source"]),
                ),
            )

    def _fill_logs_tree(self, logs):
        """把操作日志写入首页表格。"""
        self.logs_tree.delete(*self.logs_tree.get_children())
        for item in logs:
            self.logs_tree.insert(
                "",
                "end",
                values=(
                    item["created_at"],
                    item["username"] or "-",
                    ACTION_TEXT.get(item["action"], item["action"]),
                    short_path(item["detail"], max_len=80),
                ),
            )

    def detect_image(self):
        """选择图片并启动后台检测线程。"""
        path = filedialog.askopenfilename(
            title="选择待检测图片",
            filetypes=[
                ("Image files", "*.jpg *.jpeg *.png *.bmp"),
                ("All files", "*.*"),
            ],
        )
        if not path:
            return
        self._run_task(
            lambda: self._detect_image_worker(path),
            "正在检测图片...",
            error_action="detect_image_failed",
        )

    def _detect_image_worker(self, path):
        """在线程中执行 YOLO 推理，结束后回到主线程刷新界面。"""
        detector = StudyBehaviorDetector()
        output_path, events, summary = detector.predict_image(path)
        alerts = summary.get("alert_labels", [])
        self.db.record_detection(
            self.user["id"], str(path), summary, alerts, output_path=output_path
        )
        self._record_behavior_events(events, "image", path, output_path)
        self.db.log_operation(
            self.user["id"], "detect_image", f"{Path(path).name} -> {output_path.name}"
        )
        result = {
            "source": Path(path),
            "output_path": output_path,
            "events": events,
            "summary": summary,
        }
        self.after(0, lambda: self._handle_detection_done(result))

    def _handle_detection_done(self, result):
        """展示检测结果，并允许用户把错误结果保存为回流样本。"""
        self.last_detection = result
        if self.save_sample_button is not None:
            self.save_sample_button.configure(state="normal")
        self._auto_collect_suspected_sample(
            result["source"],
            result["output_path"],
            result["events"],
            result["summary"],
        )
        alerts = result["summary"].get("alert_labels", [])
        message = (
            f"输出图片：{result['output_path']}\n"
            f"检测事件：{len(result['events'])}\n"
            f"告警类别：{', '.join(alerts) if alerts else '无'}"
        )
        messagebox.showinfo("检测完成", message)
        self.refresh_dashboard()

    def detect_batch_files(self):
        """批量选择图片和视频，图片并发处理，视频队列化处理。"""
        paths = filedialog.askopenfilenames(
            title="选择批量检测文件",
            filetypes=[
                ("Media files", "*.jpg *.jpeg *.png *.bmp *.mp4 *.avi *.mov *.mkv *.wmv"),
                ("All files", "*.*"),
            ],
        )
        if not paths:
            return
        max_workers = simpledialog.askinteger(
            "图片并发数量",
            "图片检测线程数（视频会按队列顺序处理）：",
            initialvalue=2,
            minvalue=1,
            maxvalue=8,
            parent=self,
        )
        if not max_workers:
            return
        self._run_task(
            lambda: self._batch_detect_worker(list(paths), max_workers),
            "正在执行批量检测...",
            error_action="batch_detect_failed",
        )

    def _batch_detect_worker(self, paths, max_workers):
        """执行批量检测任务，图片并发，视频顺序处理。"""
        image_paths = [path for path in paths if is_image_file(path)]
        video_paths = [path for path in paths if is_video_file(path)]
        unsupported = [path for path in paths if path not in image_paths and path not in video_paths]
        total = len(image_paths) + len(video_paths)
        completed = 0
        failures = []

        self.after(0, lambda: self.batch_status_var.set(f"批量任务开始：共 {total} 个文件。"))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(self._detect_batch_image, path): path
                for path in image_paths
            }
            for future in as_completed(futures):
                completed += 1
                path = futures[future]
                try:
                    future.result()
                except Exception as exc:
                    failures.append(f"{Path(path).name}: {exc}")
                self.after(
                    0,
                    lambda completed=completed, total=total: self.batch_status_var.set(
                        f"图片检测进度：{completed}/{total}"
                    ),
                )

        for path in video_paths:
            completed += 1
            try:
                self._detect_batch_video(path)
            except Exception as exc:
                failures.append(f"{Path(path).name}: {exc}")
            self.after(
                0,
                lambda completed=completed, total=total: self.batch_status_var.set(
                    f"批量检测进度：{completed}/{total}"
                ),
            )

        if unsupported:
            failures.extend(f"{Path(path).name}: 不支持的文件类型" for path in unsupported)
        self.db.log_operation(
            self.user["id"],
            "detect_batch",
            f"files={len(paths)}, success={total - len(failures)}, failed={len(failures)}",
        )
        self.after(0, lambda failures=failures: self._show_batch_result(failures))

    def _detect_batch_image(self, path):
        """批量检测单张图片并写入检测记录。"""
        detector = StudyBehaviorDetector()
        output_path, events, summary = detector.predict_image(path)
        alerts = summary.get("alert_labels", [])
        self.db.record_detection(
            self.user["id"], str(path), summary, alerts, output_path=output_path
        )
        self._record_behavior_events(events, "image", path, output_path)
        self._auto_collect_suspected_sample(Path(path), output_path, events, summary)

    def _detect_batch_video(self, path):
        """批量检测单个视频并写入汇总记录。"""
        detector = StudyBehaviorDetector()
        output_path, summary = detector.run_video(path)
        alerts = summary.get("alert_labels", [])
        self.db.record_detection(
            self.user["id"], str(path), summary, alerts, output_path=output_path
        )
        self._record_behavior_events(
            [
                {"label": label, "confidence": 0.0, "alert": True, "reason": "视频汇总告警。"}
                for label in alerts
            ],
            "video",
            path,
            output_path,
        )

    def _show_batch_result(self, failures):
        """展示批量检测结果，失败列表不会中断已完成任务。"""
        if failures:
            message = "批量检测完成，但以下文件失败：\n" + "\n".join(failures[:12])
            if len(failures) > 12:
                message += f"\n... 还有 {len(failures) - 12} 项"
            messagebox.showwarning("批量检测完成", message)
        else:
            messagebox.showinfo("批量检测完成", "所有文件已完成检测。")
        self.batch_status_var.set("批量任务完成。")
        self.refresh_dashboard()

    def start_camera(self):
        """启动摄像头或视频流检测，防止重复打开同一个任务。"""
        if self.running:
            messagebox.showinfo("摄像头检测", "摄像头检测已经在运行。")
            return
        source = simpledialog.askstring(
            "摄像头来源",
            "请输入摄像头编号、视频路径、RTSP URL 或 HTTP 视频流 URL：",
            initialvalue="0",
            parent=self,
        )
        if source is None:
            return
        source_value = self._parse_camera_source(source)
        self.running = True
        self.camera_session = self._new_camera_session(source_value, source_type="camera")
        self._attach_study_session(self.camera_session, f"摄像头检测：{source_value}")
        self._clear_camera_queue()
        self.camera_stop_event = threading.Event()
        self._set_camera_controls(running=True)
        self._run_task(
            lambda: self._camera_worker(source_value, self.camera_stop_event),
            "摄像头检测运行中...",
            on_done=self._camera_finished,
            error_action="camera_failed",
        )

    def _camera_worker(self, source, stop_event):
        """在线程中读取摄像头并推理，画面通过队列嵌入主界面显示。"""
        try:
            import cv2
        except ImportError as exc:
            raise RuntimeError("摄像头检测需要安装 opencv-python。") from exc

        self.db.log_operation(self.user["id"], "start_camera", str(source))
        detector = StudyBehaviorDetector()
        latest_frame = None
        try:
            cap = detector._open_capture(source)
            if not cap.isOpened():
                raise RuntimeError(f"Unable to open camera/video source: {source}")

            try:
                while cap.isOpened() and not stop_event.is_set():
                    ok, frame = cap.read()
                    if not ok:
                        break
                    started = perf_counter()
                    annotated, events, summary = detector.predict_frame(frame)
                    elapsed_ms = (perf_counter() - started) * 1000
                    latest_frame = annotated
                    self._update_camera_session(events, summary, elapsed_ms)
                    self._push_camera_frame(annotated, events, summary, elapsed_ms)
            finally:
                cap.release()
        finally:
            if latest_frame is not None:
                self._record_camera_session(source, latest_frame)
            self.db.log_operation(self.user["id"], "stop_camera", str(source))

    def stop_camera(self):
        """通知摄像头线程退出，真正释放资源由检测线程完成。"""
        if not self.running or self.camera_stop_event is None:
            messagebox.showinfo("摄像头检测", "当前没有正在运行的摄像头检测。")
            return
        self.status_var.set("正在停止摄像头...")
        self.camera_stop_event.set()

    def _camera_finished(self):
        """摄像头线程结束后恢复按钮状态。"""
        self.running = False
        self.camera_stop_event = None
        self._set_camera_controls(running=False)
        self._clear_camera_queue()
        self._reset_live_preview("摄像头已停止。")

    def _set_camera_controls(self, running):
        """根据摄像头运行状态启用或禁用相关按钮。"""
        if self.camera_start_button is not None:
            self.camera_start_button.configure(state="disabled" if running else "normal")
        if self.camera_stop_button is not None:
            self.camera_stop_button.configure(state="normal" if running else "disabled")
        self._set_live_sample_button_enabled(self.camera_latest_frame is not None)
        self.camera_state_var.set("摄像头：运行中" if running else "摄像头：未运行")

    def _set_live_sample_button_enabled(self, enabled):
        """实时画面存在时允许保存误报帧，避免保存空帧。"""
        if self.live_sample_button is not None:
            self.live_sample_button.configure(state="normal" if enabled else "disabled")

    def _set_screen_controls(self, running):
        """根据屏幕识别运行状态启用或禁用相关按钮。"""
        if self.screen_start_button is not None:
            self.screen_start_button.configure(state="disabled" if running else "normal")
        if self.screen_stop_button is not None:
            self.screen_stop_button.configure(state="normal" if running else "disabled")

    def start_screen_detection(self):
        """启动屏幕识别，用于线上课堂、会议窗口或网页视频检测。"""
        if self.screen_running:
            messagebox.showinfo("屏幕识别", "屏幕识别已经在运行。")
            return
        region = self._select_screen_region_by_drag()
        if region == "cancelled":
            return
        self.screen_running = True
        self.screen_stop_event = threading.Event()
        self.screen_session = self._new_camera_session("screen", source_type="screen")
        self._attach_study_session(self.screen_session, "屏幕识别")
        self._clear_camera_queue()
        self._set_screen_controls(running=True)
        self._show_screen_control_window()
        self.winfo_toplevel().iconify()
        self._run_task(
            lambda: self._screen_worker(region, self.screen_stop_event),
            "屏幕识别运行中...",
            on_done=self._screen_finished,
            error_action="screen_detection_failed",
        )

    def stop_screen_detection(self):
        """停止屏幕识别线程。"""
        if not self.screen_running or self.screen_stop_event is None:
            messagebox.showinfo("屏幕识别", "当前没有正在运行的屏幕识别。")
            return
        self.status_var.set("正在停止屏幕识别...")
        self.screen_stop_event.set()

    def _screen_finished(self):
        """屏幕识别结束后恢复按钮和悬浮窗状态。"""
        self.screen_running = False
        self.screen_stop_event = None
        self._set_screen_controls(running=False)
        if self.screen_control_window and self.screen_control_window.winfo_exists():
            self.screen_control_window.destroy()
        self.screen_control_window = None
        self.winfo_toplevel().deiconify()
        self._clear_camera_queue()
        self._reset_live_preview("屏幕识别已停止。")

    def _screen_worker(self, region, stop_event):
        """采集屏幕画面并进行 YOLO 检测，结果嵌入主界面预览区。"""
        try:
            import cv2
            import mss
            import numpy as np
        except ImportError as exc:
            raise RuntimeError("屏幕识别需要安装 mss、numpy 和 opencv-python。") from exc

        detector = StudyBehaviorDetector()
        latest_frame = None
        self.db.log_operation(self.user["id"], "start_screen", json.dumps(region, ensure_ascii=False))
        try:
            with mss.mss() as screen_capture:
                monitor = region or screen_capture.monitors[1]
                while not stop_event.is_set():
                    started = perf_counter()
                    raw = screen_capture.grab(monitor)
                    frame = np.array(raw)
                    frame = cv2.cvtColor(frame, cv2.COLOR_BGRA2BGR)
                    annotated, events, summary = detector.predict_frame(frame)
                    elapsed_ms = (perf_counter() - started) * 1000
                    latest_frame = annotated
                    self._update_screen_session(events, summary, elapsed_ms)
                    self._push_camera_frame(annotated, events, summary, elapsed_ms)
        finally:
            if latest_frame is not None:
                self._record_screen_session(region, latest_frame)
            self.db.log_operation(self.user["id"], "stop_screen", "screen detection stopped")

    def _select_screen_region_by_drag(self):
        """用全屏截图和鼠标拖拽选择屏幕识别区域。"""
        try:
            import mss
            from PIL import Image, ImageTk
        except ImportError as exc:
            raise RuntimeError("框选屏幕区域需要安装 mss 和 Pillow。") from exc

        with mss.mss() as screen_capture:
            monitor = screen_capture.monitors[1]
            shot = screen_capture.grab(monitor)
            image = Image.frombytes("RGB", shot.size, shot.rgb)

        selector = tk.Toplevel(self)
        selector.title("拖拽选择屏幕检测区域")
        selector.attributes("-fullscreen", True)
        selector.attributes("-topmost", True)
        selector.configure(bg="#000000")
        screen_width = selector.winfo_screenwidth()
        screen_height = selector.winfo_screenheight()
        display = image.copy()
        display.thumbnail((screen_width, screen_height))
        scale_x = image.width / display.width
        scale_y = image.height / display.height
        photo = ImageTk.PhotoImage(display)
        result = {"value": None, "start": None, "rect": None}

        canvas = tk.Canvas(selector, width=display.width, height=display.height, highlightthickness=0)
        canvas.pack(expand=True)
        canvas.create_image(0, 0, image=photo, anchor="nw")
        canvas.create_text(
            18,
            18,
            anchor="nw",
            fill="#ffffff",
            text="拖拽选择识别区域；按 Enter 使用全屏；按 Esc 取消",
            font=("Microsoft YaHei UI", 14, "bold"),
        )

        def on_down(event):
            result["start"] = (event.x, event.y)
            if result["rect"]:
                canvas.delete(result["rect"])
            result["rect"] = canvas.create_rectangle(event.x, event.y, event.x, event.y, outline="#22c55e", width=3)

        def on_drag(event):
            if result["start"] and result["rect"]:
                x0, y0 = result["start"]
                canvas.coords(result["rect"], x0, y0, event.x, event.y)

        def on_up(event):
            if not result["start"]:
                return
            x0, y0 = result["start"]
            x1, y1 = event.x, event.y
            left, right = sorted((x0, x1))
            top, bottom = sorted((y0, y1))
            if right - left < 20 or bottom - top < 20:
                return
            result["value"] = {
                "left": monitor["left"] + int(left * scale_x),
                "top": monitor["top"] + int(top * scale_y),
                "width": int((right - left) * scale_x),
                "height": int((bottom - top) * scale_y),
            }
            selector.destroy()

        def use_fullscreen(_event=None):
            result["value"] = None
            selector.destroy()

        def cancel(_event=None):
            result["value"] = "cancelled"
            selector.destroy()

        canvas.bind("<ButtonPress-1>", on_down)
        canvas.bind("<B1-Motion>", on_drag)
        canvas.bind("<ButtonRelease-1>", on_up)
        selector.bind("<Return>", use_fullscreen)
        selector.bind("<Escape>", cancel)
        selector.protocol("WM_DELETE_WINDOW", cancel)
        selector.photo_ref = photo
        self.wait_window(selector)
        return result["value"]

    def _show_screen_control_window(self):
        """创建屏幕识别的置顶小控制窗，方便运行时快速调控。"""
        if self.screen_control_window and self.screen_control_window.winfo_exists():
            return
        window = tk.Toplevel(self)
        window.title("屏幕识别控制")
        window.geometry("320x210")
        window.configure(bg=COLORS["card"])
        window.attributes("-topmost", True)
        window.resizable(False, False)
        window.protocol("WM_DELETE_WINDOW", self.stop_screen_detection)
        tk.Label(
            window,
            text="屏幕识别运行中",
            bg=COLORS["card"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 12, "bold"),
        ).pack(anchor="w", padx=14, pady=(14, 6))
        tk.Label(
            window,
            text="可随时停止或保存当前帧为疑似误报样本。",
            bg=COLORS["card"],
            fg=COLORS["muted"],
            font=FONTS["small"],
            wraplength=220,
            justify="left",
        ).pack(anchor="w", padx=14)
        tk.Label(
            window,
            textvariable=self.screen_control_status_var,
            bg=COLORS["card"],
            fg=COLORS["accent"],
            font=FONTS["body"],
            wraplength=280,
            justify="left",
        ).pack(anchor="w", padx=14, pady=(10, 0))
        buttons = tk.Frame(window, bg=COLORS["card"])
        buttons.pack(fill="x", padx=14, pady=14)
        ttk.Button(
            buttons,
            text="保存当前帧",
            style="Secondary.TButton",
            command=lambda: self._save_current_live_frame("screen_manual"),
        ).pack(side="left", expand=True, fill="x", padx=(0, 8))
        ttk.Button(
            buttons,
            text="停止",
            style="Danger.TButton",
            command=self.stop_screen_detection,
        ).pack(side="left", expand=True, fill="x")
        self.screen_control_window = window

    def _update_screen_session(self, events, summary, elapsed_ms):
        """累积屏幕识别统计，复用摄像头会话字段结构。"""
        if self.screen_session is None:
            return
        current_camera = self.camera_session
        self.camera_session = self.screen_session
        try:
            self._update_camera_session(events, summary, elapsed_ms)
            self.screen_session = self.camera_session
        finally:
            self.camera_session = current_camera

    def _new_camera_session(self, source, source_type="camera"):
        """初始化一次摄像头检测会话的前端统计数据。"""
        return {
            "source": str(source),
            "source_type": source_type,
            "started_at": datetime.now(),
            "frame_count": 0,
            "event_count": 0,
            "alerts": set(),
            "label_counts": {},
            "confidence_sum": 0.0,
            "confidence_count": 0,
            "last_summary": {},
            "last_events": [],
            "avg_inference_ms": 0.0,
            "study_session_id": None,
            "behavior_samples": [],
        }

    def _update_camera_session(self, events, summary, elapsed_ms):
        """累积摄像头帧检测统计，供界面展示和最终记录入库。"""
        if not self.camera_session:
            return
        session = self.camera_session
        session["frame_count"] += 1
        session["event_count"] += len(events)
        session["last_summary"] = summary
        session["last_events"] = events
        for label in summary.get("alert_labels", []):
            session["alerts"].add(label)
        for event in events:
            label = event.get("label", "-")
            session["label_counts"][label] = session["label_counts"].get(label, 0) + 1
            session["confidence_sum"] += float(event.get("confidence", 0))
            session["confidence_count"] += 1
            if len(session["behavior_samples"]) < 500:
                sample = dict(event)
                sample["session_id"] = session.get("study_session_id")
                session["behavior_samples"].append(sample)
        if not events and session["frame_count"] % 30 == 0 and len(session["behavior_samples"]) < 500:
            session["behavior_samples"].append(
                {
                    "label": "normal",
                    "confidence": 1.0,
                    "alert": False,
                    "reason": "实时检测采样帧未发现异常行为。",
                    "session_id": session.get("study_session_id"),
                }
            )
        old_avg = session["avg_inference_ms"]
        count = session["frame_count"]
        session["avg_inference_ms"] = old_avg + (elapsed_ms - old_avg) / count

    def _flush_session_behavior_samples(self, session, source_type, source_path, output_path):
        """实时检测结束后批量写入采样行为，避免每帧同步写库拖慢画面。"""
        samples = session.get("behavior_samples") or []
        if not samples:
            samples = [
                {
                    "label": "normal",
                    "confidence": 1.0,
                    "alert": False,
                    "reason": "本次会话未记录到异常行为。",
                    "session_id": session.get("study_session_id"),
                }
            ]
        self._record_behavior_events(samples, source_type, source_path, output_path)

    def _push_camera_frame(self, frame, events, summary, elapsed_ms):
        """把最新帧放进小队列，队列满时丢弃旧帧以保持界面实时。"""
        session = self.camera_session or self.screen_session
        item = {
            "frame": frame,
            "events": events,
            "summary": summary,
            "elapsed_ms": elapsed_ms,
            "session": dict(session or {}),
        }
        try:
            self.camera_queue.put_nowait(item)
        except Full:
            try:
                self.camera_queue.get_nowait()
            except Empty:
                pass
            self.camera_queue.put_nowait(item)

    def _clear_camera_queue(self):
        """启动新摄像头任务前清空旧帧，避免显示上一轮画面。"""
        while True:
            try:
                self.camera_queue.get_nowait()
            except Empty:
                break

    def _poll_camera_frames(self):
        """主线程定时取出摄像头帧并刷新 Tkinter 预览控件。"""
        try:
            while True:
                item = self.camera_queue.get_nowait()
                self.camera_latest_frame = item["frame"]
                self.camera_latest_events = item.get("events", [])
                self.camera_latest_summary = item.get("summary", {})
                self._set_live_sample_button_enabled(True)
                self._render_camera_frame(item["frame"])
                self._render_camera_stats(item)
        except Empty:
            pass
        self.after(80, self._poll_camera_frames)

    def _render_camera_frame(self, frame):
        """把 OpenCV BGR 图像转换成 Tkinter 可显示的 PhotoImage。"""
        if self.camera_preview_label is None:
            return
        try:
            import cv2
            from PIL import Image, ImageTk
        except ImportError:
            self.camera_preview_label.configure(text="缺少 Pillow 或 OpenCV，无法显示内嵌画面。")
            return

        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        image = Image.fromarray(rgb)
        width = max(320, self.camera_preview_label.winfo_width() or 520)
        height = max(180, self.camera_preview_label.winfo_height() or 260)
        image.thumbnail((width, height))
        self.camera_image_ref = ImageTk.PhotoImage(image)
        self.camera_preview_label.configure(image=self.camera_image_ref, text="")

    def _render_camera_stats(self, item):
        """刷新摄像头检测统计文本，让用户看到实时检测结果。"""
        session = item.get("session") or {}
        alerts = sorted(session.get("alerts") or [])
        label_counts = session.get("label_counts") or {}
        avg_conf = 0.0
        if session.get("confidence_count"):
            avg_conf = session["confidence_sum"] / session["confidence_count"]
        labels_text = ", ".join(f"{key}:{value}" for key, value in label_counts.items()) or "无"
        current_labels = labels_text_from_events(item.get("events") or [])
        current_fps = 1000 / item.get("elapsed_ms", 1) if item.get("elapsed_ms", 0) else 0
        avg_fps = (
            1000 / session.get("avg_inference_ms", 1)
            if session.get("avg_inference_ms", 0)
            else 0
        )
        self.camera_stats_var.set(
            f"当前帧率：{current_fps:.1f} FPS\n"
            f"平均帧率：{avg_fps:.1f} FPS\n"
            f"当前结果：{current_labels}\n"
            f"告警：{', '.join(alerts) if alerts else '无'}\n"
            f"类别计数：{labels_text}\n"
            f"平均置信度：{avg_conf:.3f}"
        )
        if self.screen_running:
            self.screen_control_status_var.set(
                f"当前结果：{current_labels}\n"
                f"运行帧率：{current_fps:.1f} FPS\n"
                f"告警：{', '.join(alerts) if alerts else '无'}"
            )

    def _reset_live_preview(self, text):
        """停止实时检测后清空画面，避免保留最后一帧造成误解。"""
        self.camera_latest_frame = None
        self.camera_latest_events = []
        self.camera_latest_summary = {}
        self.camera_image_ref = None
        self._set_live_sample_button_enabled(False)
        if self.camera_preview_label is not None:
            self.camera_preview_label.configure(image="", text=text, bg="#0f172a", fg="#cbd5e1")
        self.camera_stats_var.set(text)

    def _record_camera_session(self, source, latest_frame):
        """摄像头停止后生成一条检测记录，供记录表和图表统计使用。"""
        try:
            import cv2
        except ImportError:
            cv2 = None

        if not self.camera_session:
            return
        session = self.camera_session
        finished_at = datetime.now()
        duration = (finished_at - session["started_at"]).total_seconds()
        alerts = sorted(session["alerts"])
        avg_conf = (
            session["confidence_sum"] / session["confidence_count"]
            if session["confidence_count"]
            else 0.0
        )

        output_path = None
        if cv2 is not None:
            output_dir = BASE_DIR / "outputs"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"camera_{current_stamp()}_last_frame.jpg"
            cv2.imwrite(str(output_path), latest_frame)

        summary = {
            "source_type": "camera",
            "frame_count": session["frame_count"],
            "event_count": session["event_count"],
            "alert_labels": alerts,
            "label_counts": session["label_counts"],
            "avg_confidence": round(avg_conf, 4),
            "avg_inference_ms": round(session["avg_inference_ms"], 4),
            "duration_seconds": round(duration, 2),
            "timestamp": finished_at.strftime("%Y-%m-%d %H:%M:%S"),
        }
        self.db.record_detection(
            self.user["id"],
            f"camera:{source}",
            summary,
            alerts,
            output_path=output_path,
        )
        self._flush_session_behavior_samples(session, "camera", f"camera:{source}", output_path)
        self._finish_study_session(session, notes=f"摄像头检测结束，来源：{source}")

    def _record_screen_session(self, region, latest_frame):
        """屏幕识别停止后写入检测记录，方便后续统计和报告使用。"""
        try:
            import cv2
        except ImportError:
            cv2 = None

        if not self.screen_session:
            return
        session = self.screen_session
        finished_at = datetime.now()
        duration = (finished_at - session["started_at"]).total_seconds()
        alerts = sorted(session["alerts"])
        avg_conf = (
            session["confidence_sum"] / session["confidence_count"]
            if session["confidence_count"]
            else 0.0
        )

        output_path = None
        if cv2 is not None:
            output_dir = BASE_DIR / "outputs"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"screen_{current_stamp()}_last_frame.jpg"
            cv2.imwrite(str(output_path), latest_frame)

        summary = {
            "source_type": "screen",
            "region": region or "primary_monitor",
            "frame_count": session["frame_count"],
            "event_count": session["event_count"],
            "alert_labels": alerts,
            "label_counts": session["label_counts"],
            "avg_confidence": round(avg_conf, 4),
            "avg_inference_ms": round(session["avg_inference_ms"], 4),
            "duration_seconds": round(duration, 2),
            "timestamp": finished_at.strftime("%Y-%m-%d %H:%M:%S"),
        }
        self.db.record_detection(
            self.user["id"],
            "screen:primary" if region is None else f"screen:{region}",
            summary,
            alerts,
            output_path=output_path,
        )
        source_name = "screen:primary" if region is None else f"screen:{region}"
        self._flush_session_behavior_samples(session, "screen", source_name, output_path)
        self._finish_study_session(session, notes="屏幕识别结束")

    def _save_current_live_frame(self, reason):
        """把当前摄像头或屏幕帧保存为疑似误报样本。"""
        if self.camera_latest_frame is None:
            messagebox.showinfo("保存样本", "当前还没有可保存的实时画面。")
            return
        self._save_suspected_false_positive_sample(
            source_name=reason,
            frame=self.camera_latest_frame,
            events=self.camera_latest_events,
            summary=self.camera_latest_summary,
            reason="用户在实时识别窗口手动保存。",
        )
        messagebox.showinfo("保存完成", "当前帧已保存为疑似误报样本。")

    def check_model_url(self):
        """检查模型或数据集 URL 可达性，并记录成功或失败日志。"""
        url = simpledialog.askstring(
            "网络资源检测",
            "请输入模型、数据集或说明文档 URL：",
            parent=self,
        )
        if not url:
            return
        self._run_task(
            lambda: self._check_url_worker(url),
            "正在检查网络资源...",
            error_action="check_url_failed",
        )

    def check_network_connectivity(self):
        """检查当前机器到指定主机端口的基础连通性。"""
        target = simpledialog.askstring(
            "网络连通性",
            "请输入 host:port，例如 8.8.8.8:53 或 github.com:443：",
            initialvalue="github.com:443",
            parent=self,
        )
        if not target:
            return
        host, port = self._parse_host_port(target, default_port=443)
        self._run_task(
            lambda: self._connectivity_worker(host, port),
            "正在测试网络连通性...",
            error_action="test_connectivity_failed",
        )

    def _connectivity_worker(self, host, port):
        """在线程里执行 socket 连通性测试，避免阻塞界面。"""
        result = test_connectivity(host=host, port=port)
        detail = json.dumps(result, ensure_ascii=False)
        action = "test_connectivity" if result.get("reachable") else "test_connectivity_failed"
        self.db.log_operation(self.user["id"], action, f"{host}:{port} {detail}")
        self.after(0, lambda: messagebox.showinfo("网络连通性结果", detail))

    def check_url_batch(self):
        """批量检测多个 URL，输入支持逗号、空格或换行分隔。"""
        text = simpledialog.askstring(
            "批量 URL 检测",
            "请输入多个 URL，可用逗号、空格或换行分隔：",
            parent=self,
        )
        urls = self._split_url_text(text)
        if not urls:
            return
        self._run_task(
            lambda: self._check_url_batch_worker(urls),
            "正在批量检测 URL...",
            error_action="check_url_batch_failed",
        )

    def _check_url_batch_worker(self, urls):
        """调用网络模块并发检测多个 URL，并把摘要展示到主台。"""
        results = check_urls_batch(urls)
        ok_count = sum(1 for item in results.values() if item.get("ok"))
        self.db.log_operation(self.user["id"], "check_url_batch", f"{ok_count}/{len(urls)} reachable")
        text = json.dumps(results, ensure_ascii=False, indent=2)
        self.after(0, lambda: self._show_text_window("批量 URL 检测结果", text))

    def download_remote_file(self):
        """下载远程文件到本地，适合测试模型或数据集文件链接。"""
        url = simpledialog.askstring("下载文件测试", "请输入文件 URL：", parent=self)
        if not url:
            return
        target = filedialog.asksaveasfilename(title="保存下载文件")
        if not target:
            return
        self._run_task(
            lambda: self._download_file_worker(url, target),
            "正在下载文件...",
            error_action="download_file_failed",
        )

    def _download_file_worker(self, url, target):
        """调用网络模块下载文件，并登记资源状态。"""
        try:
            output = download_file(url, target)
        except DownloadError as exc:
            self.db.upsert_model_resource("download_file", url, target, "failed")
            self.db.log_operation(self.user["id"], "download_file_failed", str(exc))
            self.after(0, lambda exc=exc: self._show_error("文件下载失败", exc))
            return
        self.db.upsert_model_resource("download_file", url, output, "downloaded")
        self.db.log_operation(self.user["id"], "download_file", f"{url} -> {output}")
        self.after(0, lambda: messagebox.showinfo("下载完成", f"文件已保存到：\n{output}"))

    def download_remote_image(self):
        """下载远程图片，供网络采集和误报样本测试使用。"""
        url = simpledialog.askstring("下载图片测试", "请输入图片 URL：", parent=self)
        if not url:
            return
        target_dir = filedialog.askdirectory(title="选择图片保存目录")
        if not target_dir:
            return
        self._run_task(
            lambda: self._download_image_worker(url, target_dir),
            "正在下载图片...",
            error_action="download_image_failed",
        )

    def _download_image_worker(self, url, target_dir):
        """调用网络模块下载图片，并把结果写入操作日志。"""
        try:
            output = download_image_from_url(url, target_dir)
        except DownloadError as exc:
            self.db.log_operation(self.user["id"], "download_image_failed", str(exc))
            self.after(0, lambda exc=exc: self._show_error("图片下载失败", exc))
            return
        self.db.log_operation(self.user["id"], "download_image", f"{url} -> {output}")
        self.after(0, lambda: messagebox.showinfo("下载完成", f"图片已保存到：\n{output}"))

    @staticmethod
    def _parse_host_port(text, default_port=443):
        """解析 host:port 输入，端口缺失时使用默认值。"""
        value = (text or "").strip()
        if ":" in value:
            host, port_text = value.rsplit(":", 1)
            try:
                return host.strip(), int(port_text)
            except ValueError:
                return host.strip(), default_port
        return value, default_port

    @staticmethod
    def _split_url_text(text):
        """把用户输入的多个 URL 拆成列表，兼容逗号、空格和换行。"""
        if not text:
            return []
        normalized = text.replace(",", "\n").replace(" ", "\n")
        return [item.strip() for item in normalized.splitlines() if item.strip()]

    def crawl_training_images(self):
        """Run image crawler for training images."""
        output_dir = filedialog.askdirectory(title="\u9009\u62e9\u722c\u53d6\u56fe\u7247\u4fdd\u5b58\u76ee\u5f55")
        if not output_dir:
            return
        engine = simpledialog.askstring(
            "\u56fe\u7247\u641c\u7d22\u5f15\u64ce",
            "\u8bf7\u8f93\u5165\u641c\u7d22\u5f15\u64ce\uff1abing \u6216 baidu",
            initialvalue="bing",
            parent=self,
        )
        if not engine:
            return
        max_count = simpledialog.askinteger(
            "\u6bcf\u7c7b\u56fe\u7247\u6570\u91cf",
            "\u8bf7\u8f93\u5165\u6bcf\u4e2a\u7c7b\u522b\u6700\u591a\u4e0b\u8f7d\u7684\u56fe\u7247\u6570\u91cf\uff1a",
            initialvalue=20,
            minvalue=1,
            maxvalue=200,
            parent=self,
        )
        if not max_count:
            return
        label = "\u91c7\u96c6\u8bad\u7ec3\u56fe\u7247"
        self._reset_tool_progress(label, "\u51c6\u5907\u91c7\u96c6...")
        self._run_task(
            lambda: self._crawl_training_images_worker(
                output_dir,
                max_count,
                engine.strip().lower(),
                lambda done, total, msg: self._post_tool_progress(label, done, total, msg),
            ),
            "\u6b63\u5728\u91c7\u96c6\u8bad\u7ec3\u56fe\u7247...",
            error_action="crawl_images_failed",
            show_global_progress=False,
            on_done=lambda: self._finish_tool_progress(label, "\u91c7\u96c6\u5b8c\u6210"),
            on_error=lambda exc: self._fail_tool_progress(label, f"\u91c7\u96c6\u5931\u8d25\uff1a{exc}"),
        )

    def _crawl_training_images_worker(self, output_dir, max_count, engine, progress=None):
        """Run crawler.crawl_all_classes in background."""
        from crawler import crawl_all_classes

        if progress:
            progress(1, 3, "\u6b63\u5728\u542f\u52a8\u56fe\u7247\u91c7\u96c6")
        result = crawl_all_classes(output_dir=output_dir, max_per_class=max_count, engine=engine)
        if progress:
            progress(3, 3, "\u91c7\u96c6\u4efb\u52a1\u5b8c\u6210")
        self.db.log_operation(self.user["id"], "crawl_images", json.dumps(result, ensure_ascii=False))
        self.after(0, lambda: self._show_text_window("\u56fe\u7247\u91c7\u96c6\u7ed3\u679c", json.dumps(result, ensure_ascii=False, indent=2)))

    def run_preprocess_pipeline(self):
        """Run preprocess pipeline for image dataset."""
        input_dir = filedialog.askdirectory(title="\u9009\u62e9\u539f\u59cb\u56fe\u7247\u76ee\u5f55")
        if not input_dir:
            return
        output_dir = filedialog.askdirectory(title="\u9009\u62e9\u9884\u5904\u7406\u8f93\u51fa\u76ee\u5f55")
        if not output_dir:
            return
        size = simpledialog.askinteger(
            "\u8f93\u51fa\u5c3a\u5bf8",
            "\u8bf7\u8f93\u5165\u6700\u957f\u8fb9\u76ee\u6807\u5c3a\u5bf8\uff1a",
            initialvalue=640,
            minvalue=128,
            maxvalue=2048,
            parent=self,
        )
        if not size:
            return
        label = "\u9884\u5904\u7406\u56fe\u7247\u96c6"
        self._reset_tool_progress(label, "\u51c6\u5907\u9884\u5904\u7406...")
        self._run_task(
            lambda: self._preprocess_pipeline_worker(
                input_dir,
                output_dir,
                size,
                lambda done, total, msg: self._post_tool_progress(label, done, total, msg),
            ),
            "\u6b63\u5728\u9884\u5904\u7406\u56fe\u7247\u96c6...",
            error_action="preprocess_failed",
            show_global_progress=False,
            on_done=lambda: self._finish_tool_progress(label, "\u9884\u5904\u7406\u5b8c\u6210"),
            on_error=lambda exc: self._fail_tool_progress(label, f"\u9884\u5904\u7406\u5931\u8d25\uff1a{exc}"),
        )

    def _preprocess_pipeline_worker(self, input_dir, output_dir, size, progress=None):
        """Run preprocess.cmd_pipeline in background."""
        from preprocess import cmd_pipeline

        total = max(1, len([item for item in Path(input_dir).rglob("*") if item.is_file()]))
        if progress:
            progress(1, total, "\u5f00\u59cb\u9884\u5904\u7406\u56fe\u7247\u96c6")
        result = cmd_pipeline(input_dir, output_dir, size=size)
        if progress:
            progress(total, total, "\u9884\u5904\u7406\u4efb\u52a1\u5b8c\u6210")
        self.db.log_operation(self.user["id"], "preprocess_dataset", f"{input_dir} -> {output_dir}")
        self.after(0, lambda: self._show_text_window("\u9884\u5904\u7406\u5b8c\u6210", json.dumps(result or {}, ensure_ascii=False, indent=2)))

    def validate_training_dataset(self):
        """Validate YOLO dataset."""
        dataset_dir = filedialog.askdirectory(title="\u9009\u62e9 YOLO \u6570\u636e\u96c6\u76ee\u5f55")
        if not dataset_dir:
            return
        yaml_path = filedialog.askopenfilename(
            title="\u9009\u62e9 data.yaml",
            filetypes=[("YAML", "*.yaml *.yml"), ("All files", "*.*")],
        )
        if not yaml_path:
            return
        label = "\u6821\u9a8c\u6570\u636e\u96c6"
        self._reset_tool_progress(label, "\u51c6\u5907\u6821\u9a8c...")
        self._run_task(
            lambda: self._validate_dataset_worker(
                dataset_dir,
                yaml_path,
                lambda done, total, msg: self._post_tool_progress(label, done, total, msg),
            ),
            "\u6b63\u5728\u6821\u9a8c\u6570\u636e\u96c6...",
            error_action="validate_dataset_failed",
            show_global_progress=False,
            on_done=lambda: self._finish_tool_progress(label, "\u6821\u9a8c\u5b8c\u6210"),
            on_error=lambda exc: self._fail_tool_progress(label, f"\u6821\u9a8c\u5931\u8d25\uff1a{exc}"),
        )

    def _validate_dataset_worker(self, dataset_dir, yaml_path, progress=None):
        """Run dataset_validator.validate_dataset in background."""
        from dataset_validator import validate_dataset

        if progress:
            progress(1, 2, "\u6b63\u5728\u68c0\u67e5\u56fe\u7247\u548c\u6807\u7b7e")
        result = validate_dataset(dataset_dir, yaml_path)
        if progress:
            progress(2, 2, "\u6821\u9a8c\u4efb\u52a1\u5b8c\u6210")
        self.db.log_operation(self.user["id"], "validate_dataset", json.dumps(result, ensure_ascii=False))
        self.after(0, lambda: self._show_text_window("\u6570\u636e\u96c6\u6821\u9a8c\u7ed3\u679c", json.dumps(result, ensure_ascii=False, indent=2)))

    def generate_dataset_preview(self):
        """Generate dataset preview image."""
        dataset_dir = filedialog.askdirectory(title="\u9009\u62e9\u6570\u636e\u96c6\u76ee\u5f55")
        if not dataset_dir:
            return
        target = filedialog.asksaveasfilename(
            title="\u4fdd\u5b58\u9884\u89c8\u56fe",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
        )
        if not target:
            return
        label = "\u751f\u6210\u6570\u636e\u9884\u89c8"
        self._reset_tool_progress(label, "\u51c6\u5907\u751f\u6210\u9884\u89c8...")
        self._run_task(
            lambda: self._generate_dataset_preview_worker(
                dataset_dir,
                target,
                lambda done, total, msg: self._post_tool_progress(label, done, total, msg),
            ),
            "\u6b63\u5728\u751f\u6210\u6570\u636e\u96c6\u9884\u89c8...",
            error_action="dataset_preview_failed",
            show_global_progress=False,
            on_done=lambda: self._finish_tool_progress(label, "\u9884\u89c8\u751f\u6210\u5b8c\u6210"),
            on_error=lambda exc: self._fail_tool_progress(label, f"\u751f\u6210\u5931\u8d25\uff1a{exc}"),
        )

    def _generate_dataset_preview_worker(self, dataset_dir, target, progress=None):
        """Run dataset_validator.generate_preview in background."""
        from dataset_validator import generate_preview

        if progress:
            progress(1, 2, "\u6b63\u5728\u8bfb\u53d6\u6570\u636e\u96c6\u6837\u672c")
        generate_preview(dataset_dir, target)
        if progress:
            progress(2, 2, "\u9884\u89c8\u56fe\u751f\u6210\u5b8c\u6210")
        output = Path(target)
        self.db.log_operation(self.user["id"], "generate_dataset_preview", str(output))
        self.after(0, lambda: messagebox.showinfo("\u9884\u89c8\u56fe\u5df2\u751f\u6210", f"\u5df2\u4fdd\u5b58\u5230\uff1a\\n{output}"))

    def _check_url_worker(self, url):
        """在线程中执行网络请求，失败信息会以友好弹窗展示。"""
        try:
            result = check_url(url)
        except NetworkError as exc:
            self.db.upsert_model_resource("remote_check", url, None, "failed")
            self.db.log_operation(self.user["id"], "check_url_failed", str(exc))
            self.after(0, lambda exc=exc: self._show_error("网络检测失败", exc))
            return

        detail = json.dumps(result, ensure_ascii=False)
        self.db.upsert_model_resource("remote_check", url, None, "reachable")
        self.db.log_operation(self.user["id"], "check_url", detail)
        self.after(0, lambda: messagebox.showinfo("网络检测成功", detail))

    def show_records(self):
        """用表格窗口展示检测记录，便于验收时查看历史输出。"""
        rows = self.db.list_detection_records(limit=300)
        columns = [
            ("id", "ID", 70),
            ("time", "时间", 150),
            ("user", "用户", 100),
            ("alerts", "告警", 140),
            ("source", "来源", 300),
            ("output", "输出", 300),
        ]
        values = [
            (
                record["id"],
                record["created_at"],
                record["username"] or "-",
                safe_alert_text(record["alerts_json"]),
                short_path(record["source"], 70),
                short_path(record["output_path"], 70),
            )
            for record in rows
        ]
        self._show_table_window("检测记录", columns, values)

    def show_database_summary(self):
        """展示本地 SQLite 数据库概览，供数据库查询入口使用。"""
        users = self.db.list_users() if self.auth.is_admin(self.user) else []
        records = self.db.list_detection_records(limit=1000)
        logs = self.db.list_operation_logs(limit=1000)
        resources = self.db.list_model_resources()
        text = (
            f"数据库文件：{self.db.db_path}\n\n"
            f"用户数量：{len(users) if users else '仅管理员可查看'}\n"
            f"检测记录：{len(records)} 条\n"
            f"操作日志：{len(logs)} 条\n"
            f"模型资源：{len(resources)} 条\n\n"
            f"误报样本：{self._count_db_table('misclassified_samples')} 条\n"
            f"训练周期：{self._count_db_table('training_cycles')} 条\n"
            f"数据回流日志：{self._count_db_table('data_reflux_log')} 条\n"
            f"学习会话：{self._count_db_table('study_sessions')} 条\n"
            f"行为明细：{self._count_db_table('study_behavior_records')} 条\n\n"
            "说明：当前系统使用本地 SQLite 数据库。数据库增强版已经支持误报样本、数据回流、"
            "训练周期、学习会话和逐条行为明细，后续图表与报告会优先使用这些结构化数据。"
        )
        self._show_text_window("数据库查询概览", text)

    def show_misclassified_samples(self):
        """展示误报样本表，验证误识别样本是否进入数据库闭环。"""
        rows = self.db.list_misclassified_samples() if self._db_supports("list_misclassified_samples") else []
        columns = [
            ("id", "ID", 70),
            ("image", "图片路径", 260),
            ("reviewed", "已审核", 80),
            ("relabel", "回流标签", 260),
            ("created", "创建时间", 160),
        ]
        values = [
            (
                row["id"],
                short_path(row.get("image_path"), 70),
                "是" if row.get("is_reviewed") else "否",
                short_path(row.get("relabel_label_path"), 70),
                row.get("created_at", "-"),
            )
            for row in rows
        ]
        self._show_table_window("误报样本", columns, values)

    def show_study_sessions(self):
        """展示学习会话表，查看摄像头和屏幕识别的持续监测记录。"""
        rows = self.db.list_study_sessions(limit=300) if self._db_supports("list_study_sessions") else []
        columns = [
            ("id", "ID", 70),
            ("user", "用户ID", 80),
            ("name", "会话", 180),
            ("start", "开始时间", 160),
            ("duration", "时长秒", 90),
            ("score", "专注分", 80),
            ("alerts", "告警", 80),
        ]
        values = [
            (
                row["id"],
                row.get("user_id", "-"),
                row.get("session_name") or "-",
                row.get("start_time", "-"),
                row.get("total_duration_seconds") or "-",
                row.get("focus_score") or "-",
                row.get("total_alerts") or 0,
            )
            for row in rows
        ]
        self._show_table_window("学习会话", columns, values)

    def show_behavior_records(self):
        """展示行为明细表，作为图表和报告统计的数据来源。"""
        rows = self.db.get_behavior_records(limit=500) if self._db_supports("get_behavior_records") else []
        columns = [
            ("id", "ID", 70),
            ("time", "时间", 160),
            ("behavior", "行为", 100),
            ("conf", "置信度", 90),
            ("alert", "告警", 70),
            ("source", "来源", 100),
            ("path", "路径", 260),
        ]
        values = [
            (
                row["id"],
                row.get("timestamp", "-"),
                row.get("behavior_type", "-"),
                f"{float(row.get('confidence') or 0):.3f}",
                "是" if row.get("is_alert") else "否",
                row.get("source_type", "-"),
                short_path(row.get("source_path"), 70),
            )
            for row in rows
        ]
        self._show_table_window("行为明细", columns, values)

    def show_training_cycles(self):
        """展示训练周期记录，为后续模型再训练闭环预留验收入口。"""
        rows = self.db.get_training_cycles() if self._db_supports("get_training_cycles") else []
        columns = [
            ("id", "ID", 70),
            ("version", "版本", 150),
            ("dataset", "基础数据集", 240),
            ("status", "状态", 100),
            ("map", "mAP", 90),
            ("model", "模型路径", 260),
            ("created", "创建时间", 160),
        ]
        values = [
            (
                row["id"],
                row.get("version", "-"),
                short_path(row.get("base_dataset"), 70),
                row.get("status", "-"),
                row.get("mAP") if row.get("mAP") is not None else "-",
                short_path(row.get("model_path"), 70),
                row.get("created_at", "-"),
            )
            for row in rows
        ]
        self._show_table_window("训练周期", columns, values)

    def show_logs(self):
        """用表格窗口展示操作日志，并提供 CSV 导出按钮。"""
        rows = self.db.list_operation_logs(limit=500)
        columns = [
            ("id", "ID", 70),
            ("time", "时间", 150),
            ("user", "用户", 100),
            ("action", "操作", 150),
            ("detail", "详情", 460),
        ]
        values = [
            (
                item["id"],
                item["created_at"],
                item["username"] or "-",
                ACTION_TEXT.get(item["action"], item["action"]),
                short_path(item["detail"], 100),
            )
            for item in rows
        ]
        self._show_table_window(
            "操作日志",
            columns,
            values,
            footer_buttons=[("导出 CSV", self.export_logs_csv)],
        )

    def export_logs_csv(self):
        """导出最近操作日志，方便实验报告或问题排查引用。"""
        target = filedialog.asksaveasfilename(
            title="导出日志 CSV",
            defaultextension=".csv",
            filetypes=[("CSV file", "*.csv")],
        )
        if not target:
            return
        rows = self.db.list_operation_logs(limit=1000)
        with open(target, "w", newline="", encoding="utf-8-sig") as file:
            writer = csv.writer(file)
            writer.writerow(["id", "created_at", "username", "action", "detail"])
            for item in rows:
                writer.writerow(
                    [
                        item["id"],
                        item["created_at"],
                        item["username"] or "-",
                        item["action"],
                        item["detail"] or "",
                    ]
                )
        self.db.log_operation(self.user["id"], "export_logs", str(target))
        messagebox.showinfo("导出完成", f"日志已导出到：\n{target}")

    def export_chart(self):
        """导出检测告警图表，支持柱状图、饼图和折线图。"""
        chart_choice = simpledialog.askstring(
            "图表类型",
            "请输入图表类型：bar=柱状图，pie=饼状图，line=折线图",
            initialvalue="bar",
            parent=self,
        )
        if chart_choice is None:
            return
        chart_type = chart_choice.strip().lower() or "bar"
        if chart_type not in {"bar", "pie", "line"}:
            messagebox.showwarning("图表类型错误", "图表类型只能是 bar、pie 或 line。")
            return
        target = filedialog.asksaveasfilename(
            title="保存统计图",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
        )
        if not target:
            return
        try:
            output = export_alert_chart(self._alert_stat_records(), target, chart_type=chart_type)
        except RuntimeError as exc:
            self.db.log_operation(self.user["id"], "export_chart_failed", str(exc))
            self._show_error("图表导出失败", exc)
            return
        self.db.log_operation(self.user["id"], "export_chart", f"{chart_type}: {output}")
        messagebox.showinfo("导出完成", f"统计图已保存到：\n{output}")
        self.refresh_dashboard()

    def show_model_compare(self):
        """展示 FP32、FP16 和 INT8 权重量化结果，并支持导出对比图。"""
        window = self._create_popup("模型精度量化与性能对比", "1180x760")
        window.columnconfigure(0, weight=1)
        window.rowconfigure(2, weight=1)
        window.rowconfigure(5, weight=1)

        title = tk.Label(
            window,
            text="FP32 / FP16 / INT8 模型量化对比",
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 18, "bold"),
        )
        title.grid(row=0, column=0, sticky="w", padx=22, pady=(18, 6))

        intro = (
            "点击“生成量化模型”后，会基于 models/best.pt 生成 FP16 半精度权重"
            "和 INT8 对称量化权重包，并比较文件大小、权重存储量和 torch.load 加载耗时。"
            "下方性能测试会选择一个图片数据集文件夹，统计平均、最快、最慢耗时以及检测结果。"
        )
        tk.Label(
            window,
            text=intro,
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=FONTS["body"],
            wraplength=980,
            justify="left",
        ).grid(row=1, column=0, sticky="ew", padx=22, pady=(0, 12))

        table_frame = ttk.Frame(window, style="Card.TFrame")
        table_frame.grid(row=2, column=0, sticky="nsew", padx=22)
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)
        tree = self._create_tree(
            table_frame,
            [
                ("precision", "精度", 110),
                ("file", "模型文件", 260),
                ("size", "文件大小", 100),
                ("tensor", "权重存储", 100),
                ("load", "加载耗时", 100),
                ("ratio", "体积比例", 100),
                ("status", "说明", 260),
            ],
        )

        note = tk.Label(
            window,
            text=self._quantization_explain_text(),
            bg=COLORS["page"],
            fg=COLORS["muted"],
            font=FONTS["small"],
            wraplength=980,
            justify="left",
        )
        note.grid(row=3, column=0, sticky="ew", padx=22, pady=(12, 0))

        benchmark_title = tk.Label(
            window,
            text="推理性能与检测结果对比",
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 13, "bold"),
        )
        benchmark_title.grid(row=4, column=0, sticky="w", padx=22, pady=(18, 8))

        benchmark_frame = ttk.Frame(window, style="Card.TFrame")
        benchmark_frame.grid(row=5, column=0, sticky="nsew", padx=22)
        benchmark_frame.columnconfigure(0, weight=1)
        benchmark_frame.rowconfigure(0, weight=1)
        benchmark_tree = self._create_tree(
            benchmark_frame,
            [
                ("precision", "模型", 90),
                ("avg", "平均耗时", 100),
                ("min", "最快耗时", 100),
                ("max", "最慢耗时", 100),
                ("boxes", "检测框", 80),
                ("labels", "类别", 220),
                ("conf", "平均置信度", 110),
                ("status", "状态", 300),
            ],
        )

        def refresh_benchmark_rows():
            self._fill_benchmark_tree(benchmark_tree)

        def refresh_rows():
            self._fill_quantization_tree(tree)

        footer = tk.Frame(window, bg=COLORS["page"])
        footer.grid(row=6, column=0, sticky="ew", padx=22, pady=18)
        footer.columnconfigure(0, weight=1)
        ttk.Button(
            footer,
            text="刷新对比",
            style="Secondary.TButton",
            command=lambda: (refresh_rows(), refresh_benchmark_rows()),
        ).grid(row=0, column=1, padx=(0, 10))
        ttk.Button(
            footer,
            text="导出量化图",
            style="Secondary.TButton",
            command=self.export_quantization_comparison_chart,
        ).grid(row=0, column=2, padx=(0, 10))
        ttk.Button(
            footer,
            text="测试推理性能",
            style="Secondary.TButton",
            command=lambda: self._start_model_benchmark(refresh_benchmark_rows),
        ).grid(row=0, column=3, padx=(0, 10))
        ttk.Button(
            footer,
            text="导出性能图",
            style="Secondary.TButton",
            command=self.export_model_benchmark_comparison_chart,
        ).grid(row=0, column=4, padx=(0, 10))
        ttk.Button(
            footer,
            text="生成 FP16 / INT8",
            style="Accent.TButton",
            command=lambda: self._run_task(
                self._generate_quantized_models_worker,
                "正在生成 FP16 / INT8 模型...",
                on_done=lambda: self._refresh_model_window(
                    window, lambda: (refresh_rows(), refresh_benchmark_rows())
                ),
                error_action="generate_quantized_models_failed",
            ),
        ).grid(row=0, column=5)

        refresh_rows()
        refresh_benchmark_rows()
        self.apply_visual_zoom()

    def _refresh_model_window(self, window, refresh_callback):
        """后台生成结束后刷新模型窗口和首页卡片。"""
        if window.winfo_exists():
            refresh_callback()
            self.apply_visual_zoom()
        self.refresh_dashboard()

    def _fill_quantization_tree(self, tree):
        """把量化报告写入表格，报告不存在时仍展示预期输出文件。"""
        tree.delete(*tree.get_children())
        for item in self._load_quantization_metrics():
            tree.insert(
                "",
                "end",
                values=(
                    item["precision"],
                    short_path(item["path"], 70),
                    f"{item['file_size_mb']:.2f} MB" if item["exists"] else "未生成",
                    f"{item['tensor_storage_mb']:.2f} MB" if item["tensor_storage_mb"] else "-",
                    f"{item['load_ms']:.2f} ms" if item["load_ms"] else "-",
                    ratio_text(item.get("file_ratio")),
                    item["status"],
                ),
            )

    def _load_quantization_metrics(self):
        """读取量化报告；没有报告时根据现有文件生成基础展示数据。"""
        if MODEL_QUANT_REPORT_PATH.exists():
            try:
                with open(MODEL_QUANT_REPORT_PATH, "r", encoding="utf-8") as file:
                    report = json.load(file)
                return report.get("models", [])
            except (OSError, json.JSONDecodeError):
                pass

        base_size = file_size_mb(MODEL_PT_PATH)
        rows = []
        for precision, path, status in [
            ("FP32", MODEL_PT_PATH, "原始 PyTorch 权重"),
            ("FP16", MODEL_FP16_PATH, "点击生成后得到半精度权重"),
            ("INT8", MODEL_INT8_PATH, "点击生成后得到 int8 量化权重包"),
        ]:
            size_mb = file_size_mb(path)
            rows.append(
                {
                    "precision": precision,
                    "path": str(path),
                    "exists": path.exists(),
                    "file_size_mb": size_mb,
                    "tensor_storage_mb": 0.0,
                    "load_ms": 0.0,
                    "build_ms": 0.0,
                    "file_ratio": size_mb / base_size if base_size else None,
                    "status": status,
                }
            )
        return rows

    def _quantization_explain_text(self):
        """说明 FP16 和 INT8 的差异，避免把量化误解为参数个数减少。"""
        return (
            "区别说明：FP16 是把浮点权重从 32 位改为 16 位，通常能明显降低权重存储和内存占用；"
            "INT8 是把浮点权重映射到 8 位整数并记录 scale，体积更小但需要对应的反量化或推理后端支持。"
            "两者通常不减少参数个数，减少的是每个参数占用的字节数。"
        )

    def export_quantization_comparison_chart(self):
        """把 FP32/FP16/INT8 的量化指标导出为图表。"""
        metrics = [item for item in self._load_quantization_metrics() if item["exists"]]
        if len(metrics) < 2:
            messagebox.showinfo("暂无图表数据", "请先生成 FP16 / INT8 模型，再导出量化对比图。")
            return
        target = filedialog.asksaveasfilename(
            title="保存量化对比图",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
        )
        if not target:
            return
        try:
            output = export_quantization_chart(metrics, target)
        except RuntimeError as exc:
            self._show_error("量化图表导出失败", exc)
            return
        self.db.log_operation(self.user["id"], "export_quantization_chart", str(output))
        messagebox.showinfo("导出完成", f"量化对比图已保存到：\n{output}")

    def _start_model_benchmark(self, refresh_callback):
        """选择测试数据集目录和重复轮数，然后启动模型推理性能测试。"""
        dataset_dir = filedialog.askdirectory(title="选择模型性能测试图片数据集目录")
        if not dataset_dir:
            return
        repeat = simpledialog.askinteger(
            "重复轮数",
            "每个模型遍历数据集多少轮？",
            initialvalue=1,
            minvalue=1,
            maxvalue=10,
            parent=self,
        )
        if not repeat:
            return
        self._run_task(
            lambda: self._benchmark_models_worker(dataset_dir, repeat),
            "正在测试模型推理性能...",
            on_done=refresh_callback,
            error_action="model_benchmark_failed",
        )

    def _benchmark_models_worker(self, dataset_dir, repeat):
        """用同一图片数据集测试不同精度模型的推理耗时和检测输出。"""
        dataset_dir = Path(dataset_dir)
        if not dataset_dir.exists():
            raise FileNotFoundError(f"测试数据集不存在：{dataset_dir}")
        image_paths = [
            path for path in dataset_dir.rglob("*")
            if path.is_file() and is_image_file(path)
        ]
        if not image_paths:
            raise RuntimeError("测试数据集目录中没有 jpg/jpeg/png/bmp 图片。")
        try:
            from ultralytics import YOLO
        except ImportError as exc:
            raise RuntimeError("模型性能测试需要安装 ultralytics。") from exc

        candidates = [
            ("FP32", MODEL_PT_PATH, "原始模型"),
            ("FP16", MODEL_FP16_PATH, "半精度模型"),
            ("INT8", MODEL_INT8_PATH, "INT8 量化包"),
        ]
        rows = []
        for precision, path, note in candidates:
            rows.append(
                self._benchmark_single_model(
                    YOLO, precision, path, image_paths, repeat, note
                )
            )

        report = {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dataset_dir": str(dataset_dir),
            "image_count": len(image_paths),
            "repeat": repeat,
            "models": rows,
        }
        MODELS_DIR.mkdir(parents=True, exist_ok=True)
        with open(MODEL_BENCHMARK_REPORT_PATH, "w", encoding="utf-8") as file:
            json.dump(report, file, ensure_ascii=False, indent=2)
        self.model_benchmark_metrics = rows
        self.db.log_operation(
            self.user["id"],
            "model_benchmark",
            f"{dataset_dir.name}, images={len(image_paths)}, repeat={repeat}",
        )

    def _benchmark_single_model(self, yolo_class, precision, path, image_paths, repeat, note):
        """测试单个模型文件，失败时把原因写入表格而不是中断整个对比。"""
        path = Path(path)
        if not path.exists():
            return self._benchmark_error_row(precision, path, f"{note}未生成")
        try:
            load_start = perf_counter()
            model, status_note = self._load_benchmark_model(yolo_class, precision, path)
            load_ms = (perf_counter() - load_start) * 1000
            durations = []
            aggregate = {"box_count": 0, "labels": set(), "confidences": []}
            for _round in range(repeat):
                for image_path in image_paths:
                    started = perf_counter()
                    results = model.predict(str(image_path), verbose=False)
                    durations.append((perf_counter() - started) * 1000)
                    self._accumulate_benchmark_detection_info(
                        results[0] if results else None, aggregate
                    )
            detection_info = self._finish_benchmark_detection_info(aggregate)
            return {
                "precision": precision,
                "path": str(path),
                "status": "ok",
                "status_text": status_note,
                "load_ms": round(load_ms, 4),
                "avg_ms": round(mean_value(durations), 4),
                "min_ms": round(min(durations), 4),
                "max_ms": round(max(durations), 4),
                "box_count": detection_info["box_count"],
                "labels": detection_info["labels"],
                "avg_confidence": detection_info["avg_confidence"],
            }
        except Exception as exc:
            return self._benchmark_error_row(precision, path, f"无法直接推理：{exc}")

    def _load_benchmark_model(self, yolo_class, precision, path):
        """加载 benchmark 模型；INT8 量化包会反量化后注入原模型结构。"""
        if precision != "INT8":
            return yolo_class(str(path)), "数据集测试完成"
        try:
            import torch
        except ImportError as exc:
            raise RuntimeError("INT8 测试需要安装 PyTorch。") from exc
        package = self._torch_load_checkpoint(torch, path)
        if not isinstance(package, dict) or "state_dict" not in package:
            return yolo_class(str(path)), "INT8 文件按普通模型测试完成"
        model = yolo_class(str(MODEL_PT_PATH))
        state_dict = {}
        for name, item in package["state_dict"].items():
            if not isinstance(item, dict) or "data" not in item:
                continue
            data = item["data"].float()
            scale = item.get("scale")
            state_dict[name] = data * float(scale) if scale else data
        load_result = model.model.load_state_dict(state_dict, strict=False)
        missing = len(getattr(load_result, "missing_keys", []))
        unexpected = len(getattr(load_result, "unexpected_keys", []))
        return model, f"INT8 量化包已反量化测试（missing={missing}, unexpected={unexpected}）"

    @staticmethod
    def _benchmark_error_row(precision, path, reason):
        """构造不可测试模型的表格行。"""
        return {
            "precision": precision,
            "path": str(path),
            "status": "failed",
            "status_text": reason,
            "load_ms": 0.0,
            "avg_ms": 0.0,
            "min_ms": 0.0,
            "max_ms": 0.0,
            "box_count": 0,
            "labels": [],
            "avg_confidence": 0.0,
        }

    @staticmethod
    def _accumulate_benchmark_detection_info(result, aggregate):
        """累积单张图片的检测框、类别和置信度。"""
        if result is None or result.boxes is None:
            return
        names = getattr(result, "names", {}) or {}
        for box in result.boxes:
            class_id = int(box.cls[0])
            aggregate["labels"].add(names.get(class_id, str(class_id)))
            aggregate["confidences"].append(float(box.conf[0]))
            aggregate["box_count"] += 1

    @staticmethod
    def _finish_benchmark_detection_info(aggregate):
        """把数据集累计检测信息转换成表格展示结构。"""
        return {
            "box_count": aggregate["box_count"],
            "labels": sorted(aggregate["labels"]),
            "avg_confidence": round(mean_value(aggregate["confidences"]), 4),
        }

    def _fill_benchmark_tree(self, tree):
        """把模型推理性能测试报告填入页面表格。"""
        tree.delete(*tree.get_children())
        for item in self._load_benchmark_metrics():
            labels = ", ".join(item.get("labels") or []) or "-"
            tree.insert(
                "",
                "end",
                values=(
                    item["precision"],
                    f"{item['avg_ms']:.2f} ms" if item["status"] == "ok" else "-",
                    f"{item['min_ms']:.2f} ms" if item["status"] == "ok" else "-",
                    f"{item['max_ms']:.2f} ms" if item["status"] == "ok" else "-",
                    item.get("box_count", 0) if item["status"] == "ok" else "-",
                    labels,
                    f"{item.get('avg_confidence', 0):.3f}" if item["status"] == "ok" else "-",
                    item.get("status_text", item.get("status", "-")),
                ),
            )

    def _load_benchmark_metrics(self):
        """读取模型性能测试报告，没有报告时给出待测试占位行。"""
        if MODEL_BENCHMARK_REPORT_PATH.exists():
            try:
                with open(MODEL_BENCHMARK_REPORT_PATH, "r", encoding="utf-8") as file:
                    report = json.load(file)
                return report.get("models", [])
            except (OSError, json.JSONDecodeError):
                pass
        return [
            self._benchmark_error_row("FP32", MODEL_PT_PATH, "尚未测试"),
            self._benchmark_error_row("FP16", MODEL_FP16_PATH, "尚未测试"),
            self._benchmark_error_row("INT8", MODEL_INT8_PATH, "尚未测试"),
        ]

    def export_model_benchmark_comparison_chart(self):
        """导出模型推理耗时和检测结果对比图。"""
        metrics = self._load_benchmark_metrics()
        if not any(item.get("status") == "ok" for item in metrics):
            messagebox.showinfo("暂无图表数据", "请先选择测试数据集并完成模型性能测试。")
            return
        target = filedialog.asksaveasfilename(
            title="保存模型性能对比图",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
        )
        if not target:
            return
        try:
            output = export_model_benchmark_chart(metrics, target)
        except RuntimeError as exc:
            self._show_error("性能图表导出失败", exc)
            return
        self.db.log_operation(self.user["id"], "export_model_benchmark_chart", str(output))
        messagebox.showinfo("导出完成", f"模型性能对比图已保存到：\n{output}")

    def show_false_positive_rules(self):
        """说明系统如何识别疑似误报样本，以及为什么仍需人工审核。"""
        text = (
            "疑似误报样本不会被系统直接判定为真实误报，而是进入待审核目录。\n\n"
            "当前自动收集规则：\n"
            "1. 检测框置信度低于 0.55，但仍产生了类别判断；\n"
            "2. 单张图中检测框数量异常偏多，可能存在背景干扰；\n"
            "3. 实时识别时用户点击“保存当前帧”；\n"
            "4. 后续可扩展为：短时间内类别频繁变化、同一位置框剧烈抖动、不同模型结果不一致。\n\n"
            "保存内容：原图或当前帧、预测图、预测类别、置信度、触发原因、时间和用户信息。\n"
            "这些样本建议由人工确认后再加入训练集，避免模型把自己的错误预测继续学进去。"
        )
        self._show_text_window("疑似误报样本识别规则", text)

    def _auto_collect_suspected_sample(self, source_path, output_path, events, summary):
        """根据低置信度和异常数量规则自动收集疑似误报样本。"""
        if not events:
            return
        low_conf_events = [
            event for event in events if float(event.get("confidence", 0)) < 0.55
        ]
        too_many_boxes = len(events) >= 8
        if not low_conf_events and not too_many_boxes:
            return
        reasons = []
        if low_conf_events:
            reasons.append("存在低置信度检测框，可能为误报。")
        if too_many_boxes:
            reasons.append("单张图检测框数量偏多，可能存在背景干扰。")
        self._save_suspected_false_positive_sample(
            source_name=Path(source_path).stem,
            image_path=source_path,
            preview_path=output_path,
            events=events,
            summary=summary,
            reason="；".join(reasons),
        )

    def _save_suspected_false_positive_sample(
        self,
        source_name,
        events,
        summary,
        reason,
        image_path=None,
        preview_path=None,
        frame=None,
    ):
        """保存疑似误报样本，既支持文件来源，也支持实时帧来源。"""
        SUSPECTED_DIR.mkdir(parents=True, exist_ok=True)
        sample_dir = SUSPECTED_DIR / f"{current_stamp()}_{source_name}"
        sample_dir.mkdir(parents=True, exist_ok=True)

        saved_image = ""
        saved_preview = ""
        if image_path:
            image_path = Path(image_path)
            image_target = sample_dir / f"source{image_path.suffix}"
            shutil.copy2(image_path, image_target)
            saved_image = str(image_target)
        if preview_path and Path(preview_path).exists():
            preview_path = Path(preview_path)
            preview_target = sample_dir / f"prediction{preview_path.suffix}"
            shutil.copy2(preview_path, preview_target)
            saved_preview = str(preview_target)
        if frame is not None:
            try:
                import cv2
            except ImportError:
                cv2 = None
            if cv2 is not None:
                image_target = sample_dir / "live_frame.jpg"
                cv2.imwrite(str(image_target), frame)
                saved_image = str(image_target)

        metadata = {
            "reason": reason,
            "created_by": self.user["username"],
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "source_image": saved_image,
            "prediction_preview": saved_preview,
            "summary": summary,
            "events": events,
            "review_status": "pending",
        }
        if self._db_supports("insert_misclassified_sample"):
            try:
                sample_id = self.db.insert_misclassified_sample(
                    saved_image or str(sample_dir),
                    metadata,
                    self.user["id"],
                )
                metadata["database_sample_id"] = sample_id
            except Exception as exc:
                metadata["database_error"] = str(exc)
                self.db.log_operation(self.user["id"], "save_error_sample_failed", str(exc))
        meta_path = sample_dir / "metadata.json"
        with open(meta_path, "w", encoding="utf-8") as file:
            json.dump(metadata, file, ensure_ascii=False, indent=2)
        self.db.log_operation(
            self.user["id"],
            "save_suspected_false_positive",
            str(meta_path),
        )

    def _generate_quantized_models_worker(self):
        """生成 FP16 权重和 INT8 量化权重包，并写出 JSON 对比报告。"""
        if not MODEL_PT_PATH.exists():
            raise FileNotFoundError(f"原始模型不存在：{MODEL_PT_PATH}")
        try:
            import torch
        except ImportError as exc:
            raise RuntimeError("生成 FP16 / INT8 模型需要安装 PyTorch。") from exc

        MODELS_DIR.mkdir(parents=True, exist_ok=True)

        # FP16 直接把浮点张量转换成 half，保存为新的 PyTorch checkpoint。
        fp16_start = perf_counter()
        fp16_checkpoint = self._torch_load_checkpoint(torch, MODEL_PT_PATH)
        fp16_tensor_bytes = self._convert_checkpoint_fp16(fp16_checkpoint, torch)
        torch.save(fp16_checkpoint, MODEL_FP16_PATH)
        fp16_build_ms = (perf_counter() - fp16_start) * 1000

        # INT8 保存为量化包，记录每个浮点张量的 int8 数据和 scale。
        int8_start = perf_counter()
        int8_checkpoint = self._torch_load_checkpoint(torch, MODEL_PT_PATH)
        int8_package = self._build_int8_quant_package(int8_checkpoint, torch)
        torch.save(int8_package, MODEL_INT8_PATH)
        int8_build_ms = (perf_counter() - int8_start) * 1000

        report = self._build_quantization_report(
            torch,
            fp16_tensor_bytes,
            int8_package,
            fp16_build_ms,
            int8_build_ms,
        )
        with open(MODEL_QUANT_REPORT_PATH, "w", encoding="utf-8") as file:
            json.dump(report, file, ensure_ascii=False, indent=2)

        self.db.log_operation(
            self.user["id"],
            "generate_quantized_models",
            f"{MODEL_FP16_PATH.name}, {MODEL_INT8_PATH.name}",
        )
        self.after(
            0,
            lambda: messagebox.showinfo(
                "生成完成",
                f"已生成：\n{MODEL_FP16_PATH}\n{MODEL_INT8_PATH}\n\n对比报告：\n{MODEL_QUANT_REPORT_PATH}",
            ),
        )

    @staticmethod
    def _torch_load_checkpoint(torch_module, path):
        """兼容不同 PyTorch 版本读取 YOLO checkpoint。"""
        try:
            return torch_module.load(path, map_location="cpu", weights_only=False)
        except TypeError:
            return torch_module.load(path, map_location="cpu")

    def _convert_checkpoint_fp16(self, checkpoint, torch_module):
        """递归转换 checkpoint 中的浮点张量，返回转换后的张量存储字节数。"""
        tensor_bytes = 0

        def convert(value):
            nonlocal tensor_bytes
            if torch_module.is_tensor(value):
                if value.is_floating_point():
                    converted = value.detach().half()
                    tensor_bytes += converted.numel() * converted.element_size()
                    return converted
                tensor_bytes += value.numel() * value.element_size()
                return value
            if hasattr(value, "half") and hasattr(value, "state_dict"):
                value.half()
                tensor_bytes += self._state_dict_tensor_bytes(value.state_dict(), torch_module)
                return value
            if isinstance(value, dict):
                for key, item in list(value.items()):
                    value[key] = convert(item)
                return value
            if isinstance(value, list):
                for index, item in enumerate(value):
                    value[index] = convert(item)
                return value
            if isinstance(value, tuple):
                return tuple(convert(item) for item in value)
            return value

        convert(checkpoint)
        return tensor_bytes

    def _build_int8_quant_package(self, checkpoint, torch_module):
        """把 state_dict 中的浮点权重做 per-tensor 对称 INT8 量化。"""
        state_dict, source_key = self._extract_state_dict(checkpoint)
        quantized_state = {}
        source_tensor_bytes = 0
        quantized_tensor_bytes = 0
        quantized_tensor_count = 0

        for name, tensor in state_dict.items():
            if not torch_module.is_tensor(tensor):
                continue
            detached = tensor.detach().cpu()
            source_tensor_bytes += detached.numel() * detached.element_size()

            if detached.is_floating_point():
                max_abs = float(detached.abs().max().item()) if detached.numel() else 0.0
                scale = max_abs / 127 if max_abs > 0 else 1.0
                quantized = torch_module.clamp(
                    (detached.float() / scale).round(), -128, 127
                ).to(torch_module.int8)
                quantized_state[name] = {
                    "data": quantized,
                    "scale": scale,
                    "zero_point": 0,
                    "shape": list(detached.shape),
                    "source_dtype": str(detached.dtype),
                    "scheme": "symmetric_per_tensor",
                }
                quantized_tensor_bytes += quantized.numel() * quantized.element_size()
                quantized_tensor_count += 1
            else:
                quantized_state[name] = {
                    "data": detached.clone(),
                    "scale": None,
                    "zero_point": None,
                    "shape": list(detached.shape),
                    "source_dtype": str(detached.dtype),
                    "scheme": "unchanged_non_float",
                }
                quantized_tensor_bytes += detached.numel() * detached.element_size()

        return {
            "format": "study_monitor_int8_quantized_state_dict",
            "source_model": str(MODEL_PT_PATH),
            "source_key": source_key,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "quantized_tensor_count": quantized_tensor_count,
            "source_tensor_bytes": source_tensor_bytes,
            "quantized_tensor_bytes": quantized_tensor_bytes,
            "state_dict": quantized_state,
        }

    @staticmethod
    def _extract_state_dict(checkpoint):
        """从 YOLO checkpoint 或普通 PyTorch 模型中提取 state_dict。"""
        if hasattr(checkpoint, "state_dict"):
            return checkpoint.state_dict(), "model"
        if isinstance(checkpoint, dict):
            for key in ("ema", "model"):
                value = checkpoint.get(key)
                if hasattr(value, "state_dict"):
                    return value.state_dict(), key
            tensor_items = {key: value for key, value in checkpoint.items()}
            if tensor_items:
                return tensor_items, "checkpoint"
        raise RuntimeError("无法从 best.pt 中提取可量化的 state_dict。")

    @staticmethod
    def _state_dict_tensor_bytes(state_dict, torch_module):
        """统计 state_dict 中所有张量当前占用的存储字节数。"""
        total = 0
        for tensor in state_dict.values():
            if torch_module.is_tensor(tensor):
                total += tensor.numel() * tensor.element_size()
        return total

    def _build_quantization_report(
        self,
        torch_module,
        fp16_tensor_bytes,
        int8_package,
        fp16_build_ms,
        int8_build_ms,
    ):
        """汇总文件大小、张量存储和加载耗时，供表格和图表复用。"""
        fp32_tensor_bytes = int8_package["source_tensor_bytes"]
        int8_tensor_bytes = int8_package["quantized_tensor_bytes"]
        fp32_size = file_size_mb(MODEL_PT_PATH)

        rows = [
            {
                "precision": "FP32",
                "path": str(MODEL_PT_PATH),
                "exists": MODEL_PT_PATH.exists(),
                "file_size_mb": fp32_size,
                "tensor_storage_mb": bytes_to_mb(fp32_tensor_bytes),
                "load_ms": self._measure_torch_load_ms(torch_module, MODEL_PT_PATH),
                "build_ms": 0.0,
                "file_ratio": 1.0,
                "status": "原始 PyTorch FP32 权重",
            },
            {
                "precision": "FP16",
                "path": str(MODEL_FP16_PATH),
                "exists": MODEL_FP16_PATH.exists(),
                "file_size_mb": file_size_mb(MODEL_FP16_PATH),
                "tensor_storage_mb": bytes_to_mb(fp16_tensor_bytes),
                "load_ms": self._measure_torch_load_ms(torch_module, MODEL_FP16_PATH),
                "build_ms": fp16_build_ms,
                "file_ratio": file_size_mb(MODEL_FP16_PATH) / fp32_size if fp32_size else None,
                "status": "浮点权重已转换为 FP16",
            },
            {
                "precision": "INT8",
                "path": str(MODEL_INT8_PATH),
                "exists": MODEL_INT8_PATH.exists(),
                "file_size_mb": file_size_mb(MODEL_INT8_PATH),
                "tensor_storage_mb": bytes_to_mb(int8_tensor_bytes),
                "load_ms": self._measure_torch_load_ms(torch_module, MODEL_INT8_PATH),
                "build_ms": int8_build_ms,
                "file_ratio": file_size_mb(MODEL_INT8_PATH) / fp32_size if fp32_size else None,
                "status": "INT8 对称量化权重包",
            },
        ]
        return {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "note": self._quantization_explain_text(),
            "models": rows,
        }

    def _measure_torch_load_ms(self, torch_module, path):
        """测量 torch.load 读取模型文件的耗时，不执行任何检测推理。"""
        if not Path(path).exists():
            return 0.0
        start = perf_counter()
        self._torch_load_checkpoint(torch_module, path)
        return round((perf_counter() - start) * 1000, 4)

    def save_misclassified_sample(self):
        """保存用户认为错误的检测结果，形成后续再标注和再训练样本。"""
        if self.camera_latest_frame is not None:
            self._save_current_live_frame("live_manual")
            return
        if not self.last_detection:
            messagebox.showinfo("误识别样本", "请先完成一次图片检测，再保存误识别样本。")
            return
        note = simpledialog.askstring(
            "误识别说明",
            "请简单说明哪里识别错了，便于后续标注修正：",
            parent=self,
        )
        if note is None:
            return

        RELABEL_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
        RELABEL_LABELS_DIR.mkdir(parents=True, exist_ok=True)

        source = self.last_detection["source"]
        sample_name = f"{current_stamp()}_{source.stem}"
        image_target = RELABEL_IMAGES_DIR / f"{sample_name}{source.suffix}"
        shutil.copy2(source, image_target)

        output_path = Path(self.last_detection["output_path"])
        if output_path.exists():
            preview_target = RELABEL_IMAGES_DIR / f"{sample_name}_predicted{output_path.suffix}"
            shutil.copy2(output_path, preview_target)
        else:
            preview_target = None

        label_target = RELABEL_LABELS_DIR / f"{sample_name}.json"
        label_data = {
            "source_image": str(image_target),
            "predicted_preview": str(preview_target) if preview_target else "",
            "user_note": note,
            "created_by": self.user["username"],
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "summary": self.last_detection["summary"],
            "predictions": self.last_detection["events"],
            "corrections": [],
        }
        with open(label_target, "w", encoding="utf-8") as file:
            json.dump(label_data, file, ensure_ascii=False, indent=2)

        if self._db_supports("insert_misclassified_sample"):
            try:
                sample_id = self.db.insert_misclassified_sample(
                    str(image_target),
                    label_data,
                    self.user["id"],
                )
                if self._db_supports("update_misclassified_corrected"):
                    self.db.update_misclassified_corrected(sample_id, {"user_note": note})
                if self._db_supports("update_misclassified_relabel_paths"):
                    self.db.update_misclassified_relabel_paths(
                        sample_id,
                        str(image_target),
                        str(label_target),
                    )
            except Exception as exc:
                self.db.log_operation(self.user["id"], "save_error_sample_failed", str(exc))

        self.db.log_operation(
            self.user["id"],
            "save_error_sample",
            f"{image_target.name} -> {label_target.name}",
        )
        messagebox.showinfo(
            "样本已保存",
            f"原图和预测信息已保存到 relabel 目录：\n{label_target}",
        )

    def show_packaging_guide(self):
        """展示 PyInstaller 打包命令和交付注意事项。"""
        text = (
            "推荐打包命令：\n\n"
            "pyinstaller -F -w app.py --add-data \"models;models\"\n\n"
            "交付建议：\n"
            "1. 如果模型文件较大，可以把 models/best.pt、best_fp16.pt 或 best_int8_quantized.pt 单独放网盘，并在说明文档中写下载链接。\n"
            "2. 第一次运行会自动创建 project_data/study_monitor.db，无需手动准备数据库。\n"
            "3. OpenCV、YOLO 和 PyTorch 体积较大，课程提交时可把安装包作为可选交付物。\n"
            "4. 打包前建议先运行 python app.py，确认登录、检测、日志、导出图表都能正常工作。"
        )
        self._show_text_window("打包为安装包", text)

    def export_analysis_report(self):
        """导出 Word 和 PDF 课堂学习行为分析报告。"""
        target_dir = filedialog.askdirectory(title="选择报告导出目录")
        if not target_dir:
            return
        report_format = simpledialog.askstring(
            "报告格式",
            "请输入导出格式：word、pdf 或 both",
            initialvalue="both",
            parent=self,
        )
        if report_format is None:
            return
        report_format = report_format.strip().lower() or "both"
        if report_format not in {"word", "pdf", "both"}:
            messagebox.showwarning("格式错误", "报告格式只能是 word、pdf 或 both。")
            return
        self._run_task(
            lambda: self._export_analysis_report_worker(Path(target_dir), report_format),
            "正在导出分析报告...",
            error_action="export_report_failed",
        )

    def _export_analysis_report_worker(self, target_dir, report_format):
        """汇总检测记录、模型报告和统计图，生成双格式报告。"""
        target_dir.mkdir(parents=True, exist_ok=True)
        records = self.db.list_detection_records(limit=1000)
        logs = self.db.list_operation_logs(limit=50)
        counts = count_alert_labels(self._alert_stat_records(limit=1000))
        stamp = current_stamp()
        docx_path = target_dir / f"study_behavior_report_{stamp}.docx"
        pdf_path = target_dir / f"study_behavior_report_{stamp}.pdf"
        summary = self._build_report_summary(records, counts)
        outputs = []

        if report_format in {"word", "both"}:
            self._write_word_report(docx_path, records, logs, counts, summary)
            outputs.append(docx_path)
        if report_format in {"pdf", "both"}:
            self._write_pdf_report(pdf_path, records, counts, summary)
            outputs.append(pdf_path)
        self.db.log_operation(
            self.user["id"],
            "export_report",
            ", ".join(path.name for path in outputs),
        )
        self.after(
            0,
            lambda: messagebox.showinfo(
                "报告导出完成",
                "已导出：\n" + "\n".join(str(path) for path in outputs),
            ),
        )

    def _build_report_summary(self, records, counts):
        """生成报告摘要数据，避免 Word/PDF 重复计算。"""
        total_alerts = sum(counts.values())
        top_alert = counts.most_common(1)[0][0] if counts else "无"
        return {
            "record_count": len(records),
            "total_alerts": total_alerts,
            "top_alert": top_alert,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "narrative": self._build_report_narrative(len(records), total_alerts, top_alert, counts),
        }

    def _build_report_narrative(self, record_count, total_alerts, top_alert, counts):
        """根据检测统计拼接固定模板句子，生成报告结论段。"""
        if record_count == 0:
            return [
                "本次系统尚未产生检测记录，建议先完成图片、摄像头或屏幕识别任务后再导出正式报告。",
                "当前报告仅作为系统运行状态模板展示，不代表真实课堂行为分析结论。",
            ]
        if total_alerts == 0:
            return [
                "本次检测未统计到明显异常行为告警，整体课堂状态较为平稳。",
                "建议继续保留周期性检测，以便形成更完整的学习状态趋势分析。",
            ]
        sentences = [
            f"本次共形成 {record_count} 条检测记录，累计发现 {total_alerts} 次异常行为告警。",
            f"最高频异常类别为 {top_alert}，建议教师重点关注该类行为在课堂中的出现位置和时间段。",
        ]
        if counts.get("phone", 0) > 0:
            sentences.append("系统检测到玩手机相关行为，说明课堂注意力管理仍有优化空间。")
        if counts.get("sleep", 0) > 0:
            sentences.append("系统检测到睡觉相关行为，建议结合课程时段和学生状态进一步分析原因。")
        if counts.get("eat", 0) > 0:
            sentences.append("系统检测到吃东西相关行为，可作为课堂纪律与环境管理的辅助参考。")
        sentences.append("以上结论由检测结果自动生成，最终解释应结合教师现场观察和人工复核。")
        return sentences

    def _write_word_report(self, docx_path, records, logs, counts, summary):
        """使用 python-docx 生成 Word 分析报告。"""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("导出 Word 报告需要安装 python-docx。") from exc

        document = Document()
        document.add_heading("课堂学习行为分析报告", level=0)
        document.add_paragraph(f"生成时间：{summary['generated_at']}")
        document.add_paragraph(f"检测记录数：{summary['record_count']}")
        document.add_paragraph(f"告警总数：{summary['total_alerts']}")
        document.add_paragraph(f"最高频异常类别：{summary['top_alert']}")
        document.add_heading("自动分析结论", level=1)
        for sentence in summary["narrative"]:
            document.add_paragraph(sentence)

        document.add_heading("异常类别统计", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "类别"
        table.rows[0].cells[1].text = "次数"
        for label in ["phone", "sleep", "eat"]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(counts.get(label, 0))

        document.add_heading("最近检测记录", level=1)
        for record in records[:10]:
            document.add_paragraph(
                f"{record['created_at']} | {record['username'] or '-'} | "
                f"{safe_alert_text(record['alerts_json'])} | {record['source']}"
            )

        document.add_heading("模型量化报告摘要", level=1)
        for item in self._load_quantization_metrics():
            document.add_paragraph(
                f"{item['precision']}：文件 {item['file_size_mb']:.2f} MB，"
                f"加载 {item.get('load_ms', 0):.2f} ms，{item.get('status', '')}"
            )

        document.add_heading("模型推理性能摘要", level=1)
        for item in self._load_benchmark_metrics():
            document.add_paragraph(
                f"{item['precision']}：{item.get('status_text', '-')}, "
                f"平均耗时 {item.get('avg_ms', 0):.2f} ms，"
                f"检测框 {item.get('box_count', 0)}，类别 {', '.join(item.get('labels') or []) or '无'}"
            )

        document.add_heading("最近系统日志", level=1)
        for item in logs[:10]:
            document.add_paragraph(
                f"{item['created_at']} | {item['username'] or '-'} | "
                f"{ACTION_TEXT.get(item['action'], item['action'])} | {item['detail'] or ''}"
            )
        document.save(docx_path)

    def _write_pdf_report(self, pdf_path, records, counts, summary):
        """使用 reportlab 生成 PDF 分析报告。"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise RuntimeError("导出 PDF 报告需要安装 reportlab。") from exc

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        normal = styles["BodyText"]
        title = styles["Title"]
        normal.fontName = "STSong-Light"
        title.fontName = "STSong-Light"
        styles["Heading2"].fontName = "STSong-Light"
        story = [
            Paragraph("Study Behavior Analysis Report", title),
            Paragraph(f"Generated at: {summary['generated_at']}", normal),
            Spacer(1, 12),
            Paragraph(f"Detection records: {summary['record_count']}", normal),
            Paragraph(f"Total alerts: {summary['total_alerts']}", normal),
            Paragraph(f"Most frequent alert: {summary['top_alert']}", normal),
            Spacer(1, 12),
            Paragraph("Auto-generated Conclusions", styles["Heading2"]),
        ]
        for sentence in summary["narrative"]:
            story.append(Paragraph(sentence, normal))
        story.append(Spacer(1, 12))
        data = [["Label", "Count"]] + [[label, str(counts.get(label, 0))] for label in ["phone", "sleep", "eat"]]
        table = Table(data)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 12))
        story.append(Paragraph("Recent Records", styles["Heading2"]))
        for record in records[:12]:
            story.append(
                Paragraph(
                    f"{record['created_at']} | {safe_alert_text(record['alerts_json'])} | "
                    f"{short_path(record['source'], 80)}",
                    normal,
                )
            )
        document = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        document.build(story)

    def change_password(self):
        """修改当前用户密码，空密码或旧密码错误会直接提示。"""
        old_password = simpledialog.askstring(
            "旧密码", "请输入旧密码：", parent=self, show="*"
        )
        if old_password is None:
            return
        new_password = simpledialog.askstring(
            "新密码", "请输入新密码：", parent=self, show="*"
        )
        if not new_password:
            messagebox.showwarning("输入不完整", "新密码不能为空。")
            return
        try:
            self.auth.change_password(
                self.user["username"], old_password, new_password, operator=self.user
            )
        except AuthError as exc:
            messagebox.showerror("修改失败", str(exc))
            return
        messagebox.showinfo("修改成功", "密码已更新。")
        self.refresh_dashboard()

    def manage_users(self):
        """管理员查看用户列表，并删除非默认普通用户。"""
        try:
            users = self.auth.list_users(self.user)
        except AuthError as exc:
            messagebox.showerror("权限不足", str(exc))
            return

        window = self._create_popup("用户管理", "820x500")
        window.columnconfigure(0, weight=1)
        window.rowconfigure(1, weight=1)
        tk.Label(
            window,
            text="用户管理",
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 18, "bold"),
        ).grid(row=0, column=0, sticky="w", padx=22, pady=(18, 10))
        table_frame = ttk.Frame(window, style="Card.TFrame")
        table_frame.grid(row=1, column=0, sticky="nsew", padx=22, pady=(0, 10))
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)
        tree = self._create_tree(
            table_frame,
            [
                ("id", "ID", 70),
                ("username", "用户名", 180),
                ("role", "角色", 120),
                ("created", "创建时间", 220),
            ],
        )
        for item in users:
            tree.insert(
                "",
                "end",
                values=(item["id"], item["username"], item["role"], item["created_at"]),
            )

        def selected_username():
            """读取当前表格选中的用户名，未选择时统一提示。"""
            selection = tree.selection()
            if not selection:
                messagebox.showinfo("用户管理", "请先选择一个用户。")
                return None
            return tree.item(selection[0], "values")[1]

        def reset_selected_password():
            """管理员重置选中用户密码，对接 auth.reset_password。"""
            username = selected_username()
            if not username:
                return
            new_password = simpledialog.askstring(
                "重置密码",
                f"请输入 {username} 的新密码：",
                parent=self,
                show="*",
            )
            if not new_password:
                return
            try:
                self.auth.reset_password(username, new_password, self.user)
            except AuthError as exc:
                messagebox.showerror("重置失败", str(exc))
                return
            messagebox.showinfo("重置完成", f"{username} 的密码已更新。")
            self.refresh_dashboard()

        def change_selected_role():
            """管理员修改选中用户角色，对接 auth.change_role。"""
            username = selected_username()
            if not username:
                return
            new_role = simpledialog.askstring(
                "修改角色",
                "请输入新角色：admin 或 user：",
                initialvalue="user",
                parent=self,
            )
            if not new_role:
                return
            new_role = new_role.strip().lower()
            if new_role not in {"admin", "user"}:
                messagebox.showwarning("角色错误", "角色只能填写 admin 或 user。")
                return
            try:
                self.auth.change_role(username, new_role, self.user)
            except AuthError as exc:
                messagebox.showerror("修改失败", str(exc))
                return
            messagebox.showinfo("修改完成", f"{username} 的角色已改为 {new_role}。")
            self.manage_users()

        def delete_selected():
            username = selected_username()
            if not username:
                return
            if not messagebox.askyesno("确认删除", f"确定删除用户 {username} 吗？"):
                return
            try:
                affected = self.auth.delete_user(username, self.user)
            except AuthError as exc:
                messagebox.showerror("删除失败", str(exc))
                return
            messagebox.showinfo("删除完成", f"已删除记录数：{affected}")
            self.manage_users()
            self.refresh_dashboard()

        footer = tk.Frame(window, bg=COLORS["page"])
        footer.grid(row=2, column=0, sticky="ew", padx=22, pady=(0, 18))
        ttk.Button(
            footer,
            text="重置密码",
            style="Secondary.TButton",
            command=reset_selected_password,
        ).pack(side="right", padx=(10, 0))
        ttk.Button(
            footer,
            text="修改角色",
            style="Secondary.TButton",
            command=change_selected_role,
        ).pack(side="right", padx=(10, 0))
        ttk.Button(
            footer,
            text="删除选中用户",
            style="Danger.TButton",
            command=delete_selected,
        ).pack(side="right")
        self.apply_visual_zoom()

    def _run_task(
        self,
        target,
        status,
        on_done=None,
        error_action=None,
        show_global_progress=True,
        on_error=None,
    ):
        """?? threading.Thread ????????? GUI ???????"""
        self._begin_task(status, show_global_progress=show_global_progress)

        def runner():
            try:
                target()
            except Exception as exc:
                if error_action:
                    self.db.log_operation(self.user["id"], error_action, str(exc))
                if on_error is not None:
                    self.after(0, lambda exc=exc: on_error(exc))
                self.after(0, lambda exc=exc: self._show_error("??????", exc))
            else:
                if on_done is not None:
                    self.after(0, on_done)
            finally:
                self.after(0, lambda: self._finish_task(show_global_progress=show_global_progress))

        thread = threading.Thread(target=runner, daemon=True)
        thread.start()
        return thread

    def _begin_task(self, status, show_global_progress=True):
        """?????????????????????"""
        self.active_tasks += 1
        self.status_var.set(status)
        if "tasks" in self.stat_cards:
            self.stat_cards["tasks"].update_value(str(self.active_tasks))
        self.task_count_var.set(f"?????{self.active_tasks}")

    def _finish_task(self, show_global_progress=True):
        """??????? Ready ???????????????"""
        self.active_tasks = max(0, self.active_tasks - 1)
        self.status_var.set("Ready" if self.active_tasks == 0 else "???????...")
        if "tasks" in self.stat_cards:
            self.stat_cards["tasks"].update_value(str(self.active_tasks))
        self.task_count_var.set(f"?????{self.active_tasks}")
        self.refresh_dashboard()

    def _show_error(self, title, exc):
        """统一异常提示文案，让文件、网络和 AI 推理错误都更容易理解。"""
        message = (
            f"{title}\n\n"
            f"原因：{exc}\n\n"
            "可检查：文件路径是否存在、模型权重是否放在 models 目录、网络地址是否可访问、"
            "以及依赖包是否安装完整。"
        )
        messagebox.showerror(title, message)

    @staticmethod
    def _parse_camera_source(source):
        """把纯数字摄像头编号转成 int，其它内容按路径或 URL 处理。"""
        source = source.strip()
        return int(source) if source.isdigit() else source

    def _create_popup(self, title, geometry):
        """兼容旧调用，把原弹窗内容改为右侧主台完整页面。"""
        return self._open_main_page(title, show_title=False)

    def _show_table_window(self, title, columns, values, footer_buttons=None):
        """通用表格页面，用于检测记录、日志和用户管理等列表。"""
        window = self._create_popup(title, "980x560")
        window.columnconfigure(0, weight=1)
        window.rowconfigure(1, weight=1)
        tk.Label(
            window,
            text=title,
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 18, "bold"),
        ).grid(row=0, column=0, sticky="w", padx=22, pady=(18, 10))

        table_frame = ttk.Frame(window, style="Card.TFrame")
        table_frame.grid(row=1, column=0, sticky="nsew", padx=22)
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)
        tree = self._create_tree(table_frame, columns)
        for row in values:
            tree.insert("", "end", values=row)

        footer = tk.Frame(window, bg=COLORS["page"])
        footer.grid(row=2, column=0, sticky="ew", padx=22, pady=18)
        footer.columnconfigure(0, weight=1)
        if footer_buttons:
            for index, (label, command) in enumerate(footer_buttons):
                ttk.Button(
                    footer,
                    text=label,
                    style="Secondary.TButton",
                    command=command,
                ).grid(row=0, column=index + 1, padx=(10, 0), sticky="e")
        self.apply_visual_zoom()

    def _show_text_window(self, title, text):
        """显示长文本说明，适合数据库说明、误报规则和运行注意事项。"""
        window = self._create_popup(title, "820x520")
        window.columnconfigure(0, weight=1)
        window.rowconfigure(1, weight=1)
        tk.Label(
            window,
            text=title,
            bg=COLORS["page"],
            fg=COLORS["text"],
            font=("Microsoft YaHei UI", 18, "bold"),
        ).grid(row=0, column=0, sticky="w", padx=22, pady=(18, 10))
        text_widget = tk.Text(
            window,
            wrap="word",
            bg=COLORS["card"],
            fg=COLORS["text"],
            relief="flat",
            padx=16,
            pady=14,
            font=("Microsoft YaHei UI", 10),
        )
        text_widget.insert("1.0", text)
        text_widget.configure(state="disabled")
        text_widget.grid(row=1, column=0, sticky="nsew", padx=22, pady=(0, 18))
        self.apply_visual_zoom()


class StudyMonitorApp(tk.Tk):
    """应用根窗口，负责初始化数据库、认证服务和页面切换。"""

    def __init__(self):
        super().__init__()
        self.title("Study Behavior Monitor")
        self.geometry("1120x720")
        self.minsize(980, 620)
        self.report_callback_exception = self._handle_callback_exception
        configure_app_style(self)

        # 数据库和认证服务只创建一次，页面切换时复用同一份状态。
        self.db = Database()
        self.auth = AuthService(self.db)
        self.auth.ensure_default_admin()
        self.current_frame = None
        self.show_login()

    def _handle_callback_exception(self, exc_type, exc_value, exc_tb):
        """Log exceptions raised by Tkinter callbacks in packaged windowed builds."""
        log_path = write_runtime_error("tk_callback", exc_type, exc_value, exc_tb)
        messagebox.showerror(
            "程序运行错误",
            f"界面执行时发生错误，已写入日志：\n{log_path}",
        )

    def show_login(self):
        """显示登录页，切换页面前销毁旧 Frame。"""
        if self.current_frame:
            self.current_frame.destroy()
        self.current_frame = LoginFrame(self, self.auth, self.show_main)
        self.current_frame.pack(fill="both", expand=True)

    def show_main(self, user):
        """登录成功后显示主工作台。"""
        if self.current_frame:
            self.current_frame.destroy()
        try:
            self.current_frame = MainFrame(self, self.db, self.auth, user)
            self.current_frame.pack(fill="both", expand=True)
            self.current_frame.lift()
            self.update_idletasks()
            self.db.log_operation(user["id"], "main_loaded", "Main interface loaded.")
        except Exception as exc:
            log_path = write_runtime_error("show_main", type(exc), exc, exc.__traceback__)
            self.current_frame = LoginFrame(self, self.auth, self.show_main)
            self.current_frame.pack(fill="both", expand=True)
            messagebox.showerror(
                "主界面加载失败",
                f"登录成功，但主界面加载失败。\n错误日志：{log_path}",
            )


def main():
    """程序入口函数，供 app.py 和 PyInstaller 打包入口复用。"""
    app = StudyMonitorApp()
    app.mainloop()
